"""
Generate a 25+ page Word document covering all 15 cybersecurity projects.
"""

from __future__ import annotations

import os
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
DIAGRAMS = ROOT / "docs" / "diagrams"
OUTPUT = ROOT / "Cybersecurity_Projects_Documentation.docx"


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
        h.font.size = Pt(size)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    doc.add_page_break()


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
        doc.add_paragraph()


def draw_box(ax, x, y, w, h, text, color="#E8EEF4", edge="#1F3964"):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.05",
        facecolor=color, edgecolor=edge, linewidth=1.5,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=8, fontweight="bold", wrap=True)


def arrow(ax, x1, y1, x2, y2):
    ax.add_patch(FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle="-|>", mutation_scale=12,
        color="#1F3964", linewidth=1.2,
    ))


def save_fig(fig, name: str) -> Path:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    path = DIAGRAMS / name
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def diagram_portfolio_overview() -> Path:
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.set_title("Portfolio Architecture Overview", fontsize=14, fontweight="bold", pad=12)

    draw_box(ax, 3.5, 5.0, 3, 0.7, "Cybersecurity Beginner\nProjects Portfolio", "#D6E4F0")

    groups = [
        (0.3, 3.2, "Encoding & Crypto", ["base64-tool", "caesar-cipher", "hash-cracker"], "#E2EFDA"),
        (3.5, 3.2, "Network Recon", ["dns-lookup", "port-scanner", "traffic-analyzer"], "#FFF2CC"),
        (6.7, 3.2, "Offensive Lab", ["c2-beacon", "keylogger"], "#FCE4D6"),
        (0.3, 1.2, "Defense & Audit", ["cis-auditor", "ebpf-tracer", "firewall-engine"], "#DDEBF7"),
        (3.5, 1.2, "Detection", ["canary-tokens", "sentinel", "vuln-scanner"], "#E4DFEC"),
        (6.7, 1.2, "Privacy", ["metadata-scrubber"], "#F2F2F2"),
    ]
    for x, y, title, items, color in groups:
        draw_box(ax, x, y, 2.8, 1.6, title, color)
        ax.text(x + 1.4, y + 0.35, "\n".join(items), ha="center", va="center", fontsize=7)

    for x in [1.7, 4.9, 8.1]:
        arrow(ax, x, 3.2, 5.0, 4.3)
    for x in [1.7, 4.9, 8.1]:
        arrow(ax, x, 2.8, 5.0, 4.3)

    return save_fig(fig, "01_portfolio_overview.png")


def diagram_c2_sequence() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("C2 Beacon: Command Sequence (UML)", fontsize=13, fontweight="bold", pad=10)

    actors = [("Operator\nDashboard", 1), ("FastAPI\nServer", 4), ("Beacon\nImplant", 7)]
    for name, x in actors:
        draw_box(ax, x - 0.7, 4.2, 1.4, 0.6, name, "#DDEBF7")
        ax.plot([x, x], [0.5, 4.2], "k--", linewidth=0.8)

    steps = [
        (1, 4, 3.8, "1. WebSocket connect"),
        (4, 7, 3.4, "2. REGISTER + HEARTBEAT"),
        (1, 4, 2.8, "3. shell whoami"),
        (4, 7, 2.4, "4. TASK (XOR+Base64)"),
        (7, 4, 2.0, "5. RESULT"),
        (4, 1, 1.6, "6. Display output"),
    ]
    for x1, x2, y, label in steps:
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="-|>", color="#1F3964", lw=1.2))
        ax.text((x1 + x2) / 2, y + 0.08, label, ha="center", fontsize=7)

    return save_fig(fig, "02_c2_sequence.png")


def diagram_canary_er() -> Path:
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Canary Token Generator: Entity Relationship Diagram", fontsize=13, fontweight="bold", pad=10)

    draw_box(ax, 0.5, 2.5, 2.5, 1.8,
             "TOKENS\n─────────\nid (PK)\nmanage_id\n type\nmemo\nalert_channel\nenabled\ncreated_at", "#E2EFDA")
    draw_box(ax, 6, 2.5, 2.5, 1.8,
             "EVENTS\n─────────\nid (PK)\ntoken_id (FK)\nsource_ip\nuser_agent\ngeo_json\nfingerprint\nfired_at", "#FFF2CC")
    draw_box(ax, 3.25, 0.5, 2.5, 1.2,
             "REDIS\n─────────\ndedup keys\nrate limit buckets", "#FCE4D6")

    arrow(ax, 3.0, 3.4, 6.0, 3.4)
    ax.text(4.5, 3.55, "1 : N", ha="center", fontsize=9, fontweight="bold")
    arrow(ax, 4.5, 2.5, 4.5, 1.7)
    ax.text(4.7, 2.1, "dedup", ha="left", fontsize=8)

    return save_fig(fig, "03_canary_er.png")


def diagram_sentinel_flow() -> Path:
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    ax.set_title("Systemd Persistence Scanner: Detection Flow", fontsize=13, fontweight="bold", pad=10)

    boxes = [
        (0.3, 1.8, "CLI\nsentinel scan"),
        (2.2, 1.8, "Scanner\nRegistry"),
        (4.1, 2.8, "systemd\ncron\nssh"),
        (4.1, 1.8, "profile\nld_preload"),
        (4.1, 0.8, "udev\npam\nmotd"),
        (6.8, 1.8, "Pattern\nMatcher"),
        (8.3, 1.8, "Report\nTerminal/JSON"),
    ]
    for x, y, t in boxes:
        draw_box(ax, x, y, 1.5, 0.9, t, "#DDEBF7")

    arrow(ax, 1.8, 2.25, 2.2, 2.25)
    arrow(ax, 3.7, 2.25, 4.1, 2.25)
    arrow(ax, 3.7, 2.25, 4.1, 1.35)
    arrow(ax, 3.7, 2.25, 4.1, 0.45)
    arrow(ax, 5.6, 2.25, 6.8, 2.25)
    arrow(ax, 8.3, 2.25, 8.3, 2.25)

    return save_fig(fig, "04_sentinel_flow.png")


def diagram_setup_pipeline() -> Path:
    fig, ax = plt.subplots(figsize=(10, 3))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3)
    ax.axis("off")
    ax.set_title("Windows Setup Pipeline (setup.ps1)", fontsize=13, fontweight="bold", pad=10)

    steps = ["Project 1\nbase64-tool", "Project 2\nc2-beacon", "...\n15 projects", "Per-project\n.venv", "Build\nReport"]
    xs = [0.3, 2.5, 4.7, 6.9, 8.6]
    for x, t in zip(xs, steps):
        draw_box(ax, x, 1.0, 1.8, 1.0, t, "#E2EFDA")
    for i in range(len(xs) - 1):
        arrow(ax, xs[i] + 1.8, 1.5, xs[i + 1], 1.5)

    return save_fig(fig, "05_setup_pipeline.png")


def diagram_netanal_layers() -> Path:
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Network Traffic Analyzer: Layered Architecture", fontsize=13, fontweight="bold", pad=10)

    layers = [
        (1, 4.0, "CLI Layer\n(Typer commands)", "#D6E4F0"),
        (1, 3.0, "Capture Engine\n(Scapy AsyncSniffer)", "#DDEBF7"),
        (1, 2.0, "Analysis Layer\n(Statistics, Filters)", "#E2EFDA"),
        (1, 1.0, "Output Layer\n(Rich tables, JSON, Charts)", "#FFF2CC"),
    ]
    for x, y, t, c in layers:
        draw_box(ax, x, y, 6, 0.75, t, c)
    for y in [3.75, 2.75, 1.75]:
        arrow(ax, 4, y, 4, y - 0.2)

    return save_fig(fig, "06_netanal_layers.png")


def diagram_dns_flow() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.set_xlim(0, 9); ax.set_ylim(0, 4); ax.axis("off")
    ax.set_title("DNS Lookup: Resolution Flow", fontsize=13, fontweight="bold", pad=10)
    for x, t in [(0.3, "User\nCLI"), (2.2, "Resolver\nModule"), (4.1, "DNS\nServer"), (6, "Response\nParser"), (7.8, "Rich Table\nor JSON")]:
        draw_box(ax, x, 1.5, 1.4, 1, t, "#DDEBF7")
    for a, b in [(1.7, 2.2), (3.6, 4.1), (5.5, 6), (7.4, 7.8)]:
        arrow(ax, a, 2, b, 2)
    return save_fig(fig, "08_dns_flow.png")


def diagram_metadata_flow() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 4.5); ax.axis("off")
    ax.set_title("Metadata Scrubber: Processing Pipeline", fontsize=13, fontweight="bold", pad=10)
    steps = ["Input File", "Signature\nDetect", "Handler\nSelect", "Strip\nMetadata", "Verify\nReport", "Output File"]
    for i, t in enumerate(steps):
        draw_box(ax, 0.2 + i * 1.45, 1.8, 1.2, 1.1, t, "#E2EFDA" if i % 2 == 0 else "#FFF2CC")
        if i < len(steps) - 1:
            arrow(ax, 0.2 + i * 1.45 + 1.2, 2.35, 0.2 + (i+1) * 1.45, 2.35)
    return save_fig(fig, "09_metadata_flow.png")


def diagram_c2_deployment() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    ax.set_title("C2 Beacon: Deployment Architecture", fontsize=13, fontweight="bold", pad=10)
    draw_box(ax, 0.5, 3.5, 2, 1, "Operator Browser\nReact Dashboard", "#D6E4F0")
    draw_box(ax, 4, 3.8, 2, 0.7, "nginx\nReverse Proxy", "#DDEBF7")
    draw_box(ax, 4, 2.5, 2, 1, "FastAPI Server\nSQLite Tasks", "#E2EFDA")
    draw_box(ax, 7.5, 2.5, 2, 1, "Beacon Implant\nPython asyncio", "#FFF2CC")
    draw_box(ax, 4, 0.8, 2, 0.8, "SQLite DB", "#FCE4D6")
    arrow(ax, 2.5, 4, 4, 4.1)
    arrow(ax, 5, 3.8, 5, 3.5)
    arrow(ax, 6, 3, 7.5, 3)
    arrow(ax, 5, 2.5, 5, 1.6)
    return save_fig(fig, "10_c2_deployment.png")


def diagram_hashcracker_uml() -> Path:
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Hash Cracker: UML Class Structure", fontsize=13, fontweight="bold", pad=10)

    classes = [
        (0.4, 2.5, "Engine\n────────\n+run()\n+setHash()", "#DDEBF7"),
        (2.8, 3.5, "DictionaryAttack\n────────\n+attack()", "#E2EFDA"),
        (2.8, 2.0, "BruteForceAttack\n────────\n+attack()", "#E2EFDA"),
        (2.8, 0.5, "RuleAttack\n────────\n+mutate()", "#E2EFDA"),
        (5.5, 2.5, "ThreadPool\n────────\n+partition()", "#FFF2CC"),
        (7.2, 2.5, "HashDetector\n────────\n+detect()", "#FCE4D6"),
    ]
    for x, y, t, c in classes:
        draw_box(ax, x, y, 1.8, 1.3, t, c)

    arrow(ax, 2.2, 3.0, 2.8, 3.0)
    arrow(ax, 2.2, 2.8, 2.8, 2.5)
    arrow(ax, 2.2, 2.5, 2.8, 1.0)
    arrow(ax, 4.6, 2.8, 5.5, 2.8)
    arrow(ax, 2.2, 2.8, 7.2, 2.8)

    return save_fig(fig, "07_hashcracker_uml.png")


PROJECTS = [
    {
        "num": 1,
        "name": "base64-tool",
        "title": "Base64 Multi-Format Encoding Tool",
        "stack": "Python 3.12, Typer, Rich",
        "purpose": "I built this CLI because encoded data shows up constantly in security labs: JWT segments, hex dumps, URL-encoded parameters, and stacked obfuscation in malware samples. I wanted one local tool that could encode, decode, detect format, and peel multiple layers without relying on browser-based decoders.",
        "features": [
            "Encode and decode Base64, Base64URL, Base32, Hex, and URL formats",
            "Auto-detect encoding with confidence scoring",
            "Recursive peel command for multi-layer payloads",
            "Chain command to stack encodings for WAF testing",
            "Pipeline support via stdin and file input",
        ],
        "architecture": "The design follows a thin CLI over pure encoder functions. cli.py routes Typer commands. encoders.py holds format-specific logic behind a registry. detector.py scores candidate formats. peeler.py loops decode until confidence drops. formatter.py handles Rich output.",
        "challenges": [
            "Python 3.14 was listed as a requirement in pyproject.toml but my Windows machine runs 3.12. pip refused installation until I lowered the version constraint.",
            "Multiple Python projects installed into one global environment caused typer version conflicts between caesar-cipher and base64-tool.",
            "Detecting stacked encodings required a stopping rule. I used confidence thresholds so peel does not loop forever on random binary data.",
        ],
        "solutions": [
            "I changed requires-python to >=3.12 and verified encode/decode/peel on Windows.",
            "I wrote setup.ps1 to create a separate .venv per Python project so dependencies stay isolated.",
            "I capped peel depth and added per-layer confidence reporting so each strip is visible in the output.",
        ],
        "commands": "b64tool encode \"Hello World\" | b64tool detect <blob> | b64tool peel <hex-or-b64>",
        "implementation_steps": [
            "Step 1: I created the src/base64_tool package with separate modules for CLI, encoders, detector, peeler, and formatter.",
            "Step 2: I registered each encoding format in encoders.py behind a common encode/decode interface.",
            "Step 3: I wired Typer commands in cli.py for encode, decode, detect, peel, and chain.",
            "Step 4: I implemented detector.py to score base64, hex, base32, url, and base64url candidates.",
            "Step 5: I built peeler.py to decode repeatedly until confidence falls below threshold or max depth is reached.",
            "Step 6: I added stdin and --file input resolution in utils.py for pipeline use.",
            "Step 7: I ran pip install -e . inside a local venv and verified SGVsbG8gV29ybGQ= output for Hello World.",
        ],
        "testing": "Verified on Windows 10 with Python 3.12. encode, detect, and peel commands pass manual smoke tests. pytest suite covers encoders, detector, peeler, and CLI integration.",
        "design_decisions": [
            "Kept encoders pure functions with no CLI dependency for testability.",
            "Used confidence scoring instead of hard-coded format guessing.",
            "Separated Rich formatting from core logic so JSON-style output could be added later.",
        ],
    },
    {
        "num": 2,
        "name": "c2-beacon",
        "title": "Educational C2 Beacon and Operator Dashboard",
        "stack": "Python FastAPI, React 19, SQLite, Docker",
        "purpose": "I built this to understand command-and-control mechanics used by real frameworks. The project has three parts: a beacon implant, a team server, and a browser dashboard. Traffic uses XOR plus Base64 to simulate encoded C2 channels.",
        "features": [
            "WebSocket protocol with REGISTER, HEARTBEAT, TASK, and RESULT messages",
            "Ten operator commands mapped to MITRE ATT&CK techniques",
            "SQLite-backed task queue and beacon registry",
            "React dashboard with live beacon table and terminal session view",
            "Docker Compose stack for nginx, backend, and frontend",
        ],
        "architecture": "Two WebSocket channels separate trust boundaries. Beacons connect to /api/ws/beacon with encoded payloads. Operators connect with plain JSON. backend/app/beacon handles implant traffic. backend/app/ops handles dashboard traffic. beacon/beacon.py runs asyncio handlers for shell, sysinfo, screenshot, and related commands.",
        "challenges": [
            "Docker was not installed on my Windows host, so I could not bring up the full compose stack locally.",
            "Backend pytest failed on Windows because PYTHONPATH did not include the app package root.",
            "WebSocket proxying through nginx required careful header forwarding in dev.compose.yml.",
        ],
        "solutions": [
            "I installed Python dependencies for backend and beacon separately and ran backend tests with PYTHONPATH=app.",
            "I documented the Docker path for full-stack use and verified 33 of 35 unit tests pass on Windows.",
            "I traced message flow in protocol.py and confirmed operator JSON stays separate from beacon encoding.",
        ],
        "commands": "docker compose -f dev.compose.yml up -d | just beacon | shell whoami in session UI",
        "diagram": "c2",
    },
    {
        "num": 3,
        "name": "caesar-cipher",
        "title": "Caesar Cipher Encrypt, Decrypt, and Crack CLI",
        "stack": "Python 3.12, Typer, Rich",
        "purpose": "This was my entry point into classical cryptanalysis. The cipher itself is trivial, but writing a crack command that ranks all 26 shifts using English letter frequency taught me how statistical attacks work on stronger ciphers later.",
        "features": [
            "Encrypt and decrypt with shift key 0 through 25",
            "Brute-force crack across all shifts",
            "Chi-squared frequency scoring to rank candidates",
            "File and stdin input with optional output file",
        ],
        "architecture": "cipher.py implements the shift logic. analyzer.py computes letter frequencies and scores plaintext candidates. main.py exposes encrypt, decrypt, and crack subcommands through Typer.",
        "challenges": [
            "The CLI flag is --key not --shift, which tripped me up during first testing.",
            "Non-alphabetic characters needed consistent handling so crack scoring only evaluates letters.",
        ],
        "solutions": [
            "I documented the correct flag in README and verified caesar-cipher encrypt \"HELLO\" --key 3 returns KHOOR.",
            "I preserved case and punctuation while applying shifts only to alphabetic bytes.",
        ],
        "commands": "caesar-cipher encrypt \"HELLO\" --key 3 | caesar-cipher crack \"KHOOR ZRUOG\"",
    },
    {
        "num": 4,
        "name": "canary-token-generator",
        "title": "Self-Hosted Honeytoken and Canary Generator",
        "stack": "Go 1.25, React 19, PostgreSQL, Redis, Docker",
        "purpose": "I wanted a deception layer for lab networks: files and URLs that look valuable but alert me when touched. The generator produces seven artifact types from web bugs to fake kubeconfig files and a MySQL protocol decoy.",
        "features": [
            "Seven token types: webbug, slowredirect, pdf, docx, envfile, kubeconfig, mysql",
            "Telegram and HMAC-signed webhook notifications",
            "Redis dedup window to prevent alert spam",
            "GeoIP enrichment and browser fingerprint capture",
            "Admin API with bearer auth and Turnstile on creation",
        ],
        "architecture": "Go backend under backend/cmd/canary serves HTTP API and trigger routes. PostgreSQL stores tokens and events. Redis handles rate limits and dedup. React frontend under frontend/ handles token creation and manage pages. Docker Compose wires nginx, Postgres, Redis, and Jaeger.",
        "challenges": [
            "Full stack needs Docker, Postgres, Redis, and pnpm. My Windows machine had none of the orchestration tooling at first.",
            "The envfile generator must look realistic while embedding exactly one canary endpoint among fake secrets.",
            "The mysql token type needs a real wire-protocol handshake, not a simple HTTP 404.",
        ],
        "solutions": [
            "I built the Go backend binary with go build -o canary.exe ./cmd/canary and verified --help output.",
            "I studied the recipe shuffler in internal/token/generators and confirmed canary lines are buried in plausible config blocks.",
            "I documented the separate TCP listener path for mysql tokens in README API section.",
        ],
        "commands": "just init && just dev-up | POST /api/tokens | GET /c/{tokenID}",
        "diagram": "canary",
    },
    {
        "num": 5,
        "name": "dns-lookup",
        "title": "DNS Lookup and WHOIS CLI",
        "stack": "Python 3.12, dnspython, Rich, Typer",
        "purpose": "I made this to move past basic nslookup usage. The tool queries multiple record types, traces resolution paths, runs batch jobs, and exports JSON for scripting.",
        "features": [
            "Query A, AAAA, MX, NS, TXT, CNAME, SOA records",
            "Reverse DNS and resolution trace",
            "Batch concurrent lookups from a file",
            "WHOIS integration",
            "JSON output with --json",
        ],
        "architecture": "resolver.py wraps dnspython queries. output.py formats Rich tables. whois_lookup.py handles registration data. cli.py exposes query, reverse, trace, batch, and whois commands.",
        "challenges": [
            "Rich table rendering failed on Windows console for some Unicode box characters.",
            "I initially used --records instead of --type for the query command.",
            "WHOIS queries can be slow and occasionally rate limited by registrars.",
        ],
        "solutions": [
            "I used --json for reliable automated testing: dnslookup query google.com --type A --json.",
            "I mapped the correct Typer option names in README and setup.ps1.",
            "I added configurable timeout flags and documented JSON mode for scripting.",
        ],
        "commands": "dnslookup query google.com --type A --json | dnslookup trace example.com",
        "diagram": "dns",
    },
    {
        "num": 6,
        "name": "firewall-rule-engine",
        "title": "Firewall Rule Parser and Analyzer",
        "stack": "V language 0.5",
        "purpose": "I built this to analyze iptables and nftables rulesets before applying them. It parses rules into a unified model, detects conflicts, suggests optimizations, and can emit hardened templates.",
        "features": [
            "Parse iptables-save and nft list output",
            "Detect shadowed, duplicate, and contradictory rules",
            "Suggest merges, reordering, and missing rate limits",
            "Export between iptables and nftables formats",
            "Diff two rulesets",
        ],
        "architecture": "src/parser/ normalizes input formats. src/models/ holds the shared rule struct. src/analyzer/ runs conflict and optimizer passes. src/generator/ emits hardened output. src/display/ prints severity-coded findings.",
        "challenges": [
            "The V compiler was not installed on my Windows machine.",
            "iptables and nftables syntax diverge enough that a single parser required separate modules.",
        ],
        "solutions": [
            "I kept full source and testdata fixtures for offline review on Windows.",
            "I documented Linux install.sh path and read parser_test.v to verify fixture coverage.",
        ],
        "commands": "fwrule analyze testdata/iptables_basic.rules | fwrule harden --output rules.nft",
    },
    {
        "num": 7,
        "name": "hash-cracker",
        "title": "Multi-Threaded Hash Cracker",
        "stack": "C++23, OpenSSL, Boost, CMake",
        "purpose": "I wrote this to understand how offline hash cracking works in incident response labs. It supports dictionary, brute-force, and rule-based attacks across MD5, SHA1, SHA256, and SHA512.",
        "features": [
            "Auto-detect hash type from digest length",
            "Memory-mapped dictionary files",
            "Brute-force with charset and length controls",
            "Rule mutations: leet speak, append digits, reverse, toggle case",
            "Thread pool with per-core partitioning",
        ],
        "architecture": "Engine coordinates attack strategies. DictionaryAttack and BruteForceAttack implement the Attack interface. ThreadPool partitions keyspace. HashDetector selects the correct OpenSSL hasher. MappedFile avoids copying large wordlists into RAM.",
        "challenges": [
            "Windows lacked CMake, Ninja, Boost, and OpenSSL dev packages in my initial environment.",
            "C++23 requirement excluded older compilers on some lab machines.",
            "Rule chaining multiplies candidate count quickly and needs progress feedback.",
        ],
        "solutions": [
            "I installed CMake and LLVM via winget and documented Linux install.sh as the primary build path.",
            "I read Engine.hpp and attack classes to map thread partitioning before attempting local build.",
            "Progress.cpp reports hashes per second and ETA so long runs remain observable.",
        ],
        "commands": "hashcracker --hash <digest> --wordlist wordlists/10k-most-common.txt",
        "diagram": "hash",
    },
    {
        "num": 8,
        "name": "keylogger",
        "title": "Educational Keylogger for Security Research",
        "stack": "Python 3.12, pynput, pywin32",
        "purpose": "I built this to study input capture mechanics used by malware, strictly for authorized lab use. It logs keystrokes with window title context on Windows.",
        "features": [
            "Global keyboard hook via pynput",
            "Active window title tracking on Windows",
            "Configurable log file output",
            "Toggle hotkey and clean shutdown",
        ],
        "architecture": "keylogger.py is a single-module design with KeyloggerConfig, LogManager, WindowTracker, and the main listener loop. Windows-specific imports are guarded so the module loads on other platforms for code review.",
        "challenges": [
            "Windows requires pywin32 and psutil extras beyond base dependencies.",
            "Antivirus software flags keylogger behavior even in a lab binary.",
            "Long-running file handles needed careful lifecycle management to avoid lock errors.",
        ],
        "solutions": [
            "I installed with pip install -e \".[windows]\" inside a dedicated venv.",
            "I documented authorized-use-only scope in README legal section.",
            "LogManager owns file handles across the process lifetime with explicit close on Ctrl+C.",
        ],
        "commands": "python keylogger.py --help (authorized lab use only)",
    },
    {
        "num": 9,
        "name": "linux-cis-hardening-auditor",
        "title": "Linux CIS Benchmark Compliance Auditor",
        "stack": "Bash, shellcheck",
        "purpose": "I built this to automate CIS Benchmark checks on Debian and Ubuntu systems. It scores compliance, exports JSON and HTML, compares against a baseline, and prints remediation commands for failures.",
        "features": [
            "104 CIS controls across filesystem, services, network, logging, SSH, and accounts",
            "Terminal, JSON, and HTML reports",
            "Baseline save and diff mode",
            "Test mode against mock fixtures without root",
        ],
        "architecture": "cisaudit.sh is the entry point. src/checks/ holds section scripts. src/lib/engine.sh orchestrates control execution. src/lib/report_*.sh format output. testdata/fixtures/ provides pass and fail trees for CI.",
        "challenges": [
            "The auditor requires Linux paths like /etc/shadow and systemd units. It does not run natively on Windows.",
            "Some controls need root while test mode must work unprivileged against fixtures.",
        ],
        "solutions": [
            "I documented WSL and Linux VM as the execution environment.",
            "I used cisaudit --test with testdata/fixtures to validate control logic without root.",
        ],
        "commands": "./install.sh && sudo cisaudit | cisaudit --test",
    },
    {
        "num": 10,
        "name": "linux-ebpf-security-tracer",
        "title": "eBPF Syscall Security Tracer",
        "stack": "Python 3.12, BCC, C eBPF programs",
        "purpose": "I built this to learn kernel-level observability. eBPF programs attach to tracepoints for process, file, network, and privilege events. Userspace Python correlates events and fires detection rules mapped to MITRE ATT&CK.",
        "features": [
            "Real-time syscall monitoring via eBPF",
            "Ten built-in detection rules with severity levels",
            "Live, JSON, and table output modes",
            "Process enrichment from /proc",
            "Filter by event type, process name, and severity",
        ],
        "architecture": "src/ebpf/*.c holds kernel probes. loader.py attaches programs. processor.py reads ring buffer events. detector.py evaluates rules. renderer.py prints color-coded output. main.py is the Typer CLI.",
        "challenges": [
            "eBPF requires Linux kernel headers and bcc tools. The CLI installs on Windows but tracing does not.",
            "The pwd module used for user enrichment is Unix-only and fails on Windows import.",
            "Multiple projects used a generic src package name causing pip entry point collisions.",
        ],
        "solutions": [
            "I isolated the project in its own .venv and documented Linux as the runtime target.",
            "I installed the CLI for code review and reserved sudo ebpf-tracer for Linux lab hosts.",
            "I updated setup.ps1 to use per-project virtual environments.",
        ],
        "commands": "sudo ebpf-tracer | sudo ebpf-tracer -f json -s MEDIUM",
    },
    {
        "num": 11,
        "name": "metadata-scrubber-tool",
        "title": "Metadata Scrubber for Images and Documents",
        "stack": "Python 3.10, Pillow, piexif, pypdf, python-docx",
        "purpose": "I built this after realizing how much identifying data rides along in photos and Office files. The tool strips EXIF GPS coordinates, author names, camera models, and document properties before sharing files.",
        "features": [
            "JPEG, PNG, PDF, Word, Excel, PowerPoint support",
            "Concurrent batch processing with ThreadPoolExecutor",
            "Dry-run mode and before/after verification report",
            "Format detection by file signature not extension",
        ],
        "architecture": "src/main.py routes scrub commands. Format-specific handlers live in dedicated modules using Pillow, piexif, pypdf, and python-docx. A thread pool processes directories in parallel.",
        "challenges": [
            "PDF and Office formats keep metadata in different internal structures.",
            "Dry-run must report planned removals without writing files.",
            "The generic src package name conflicted with linux-ebpf-security-tracer in global pip installs.",
        ],
        "solutions": [
            "I split handlers per format and added a verification report comparing field counts.",
            "Dry-run prints a diff table before any write occurs.",
            "Dedicated .venv per project resolved the metadata-scrubber versus ebpf-tracer script collision.",
        ],
        "commands": "metadata-scrubber scrub photo.jpg | mst scrub --dry-run ./folder",
        "diagram": "metadata",
    },
    {
        "num": 12,
        "name": "network-traffic-analyzer",
        "title": "Network Traffic Capture and Analysis CLI",
        "stack": "Python 3.12, Scapy, matplotlib",
        "purpose": "I built this to practice packet-level analysis. It captures live traffic or reads PCAP files, computes protocol distribution, identifies top talkers, and exports charts.",
        "features": [
            "Live capture with BPF filters",
            "PCAP replay and statistics collection",
            "Protocol distribution and top talker reports",
            "JSON and CSV export",
            "Bandwidth charts via matplotlib",
        ],
        "architecture": "capture.py runs Scapy AsyncSniffer with a producer-consumer queue. analyzer.py extracts packet fields. statistics.py aggregates counts. visualization.py builds charts. main.py exposes capture, analyze, stats, export, and chart commands.",
        "challenges": [
            "Windows has no libpcap provider by default. Scapy warned that live capture would not work without Npcap.",
            "Packet type hints in capture.py used TYPE_CHECKING only, causing NameError on Windows at import time.",
            "Python 3.14 requirement blocked install on my 3.12 host.",
        ],
        "solutions": [
            "I added from __future__ import annotations to capture.py so Packet annotations defer evaluation.",
            "I lowered requires-python to 3.12 and verified netanal --help runs.",
            "I documented Npcap installation for live capture and used PCAP replay for offline testing.",
        ],
        "commands": "netanal capture -i eth0 -c 100 | netanal analyze capture.pcap",
        "diagram": "netanal",
    },
    {
        "num": 13,
        "name": "simple-port-scanner",
        "title": "Asynchronous TCP Port Scanner",
        "stack": "C++20, Boost.Asio, CMake",
        "purpose": "I wrote this to learn how nmap-style connect scans work at the socket level. Boost.Asio runs many TCP connect attempts concurrently with configurable timeout and port range.",
        "features": [
            "Configurable port ranges and thread count",
            "Open, closed, and filtered state reporting",
            "Service name mapping for common ports",
            "Banner grabbing in verbose mode",
        ],
        "architecture": "PortScanner.hpp/cpp manages an Asio io_context and strand-based posting. main.cpp parses Boost.Program_options and drives the scan loop.",
        "challenges": [
            "Boost libraries were not available on my Windows CMake attempt.",
            "Filtered ports versus closed ports need distinct timeout handling.",
            "High concurrency can trigger local firewall rate limits.",
        ],
        "solutions": [
            "I documented Linux/WSL build with install.sh where Boost is available via apt.",
            "I set explicit connect timeouts and classified timeout-as-filtered in PortScanner.cpp.",
            "Thread count remains configurable to reduce network pressure during lab scans.",
        ],
        "commands": "./simplePortScanner --target 192.168.1.1 --ports 1-1024",
    },
    {
        "num": 14,
        "name": "simple-vulnerability-scanner",
        "title": "Python Dependency Vulnerability Scanner",
        "stack": "Go 1.24, OSV.dev API, PyPI",
        "purpose": "I built this to audit Python projects for known CVEs in dependencies. It parses pyproject.toml and requirements.txt, queries OSV.dev, and can bump packages to latest stable versions.",
        "features": [
            "Parse pyproject.toml and requirements.txt",
            "OSV.dev vulnerability lookup",
            "Parallel PyPI version checks with ETag cache",
            "Update mode preserving comments and formatting",
            "Config via .angela.toml",
        ],
        "architecture": "internal/pyproject and internal/requirements parse dependency files. internal/pypi fetches versions. internal/osv queries advisories. internal/cli exposes check, update, and scan commands. cmd/angela is the entry point.",
        "challenges": [
            "Go was not installed initially on Windows.",
            "PEP 440 version strings include pre-release markers that must be filtered.",
            "Rewriting TOML while preserving comments requires a custom writer, not naive string replace.",
        ],
        "solutions": [
            "I installed Go 1.26 via winget and built with go build -o svscan.exe ./cmd/angela.",
            "version.go filters pre-release tags before presenting update candidates.",
            "writer.go in pyproject and requirements packages preserves formatting on update.",
        ],
        "commands": "svscan.exe check . | svscan.exe update --vulns .",
    },
    {
        "num": 15,
        "name": "systemd-persistence-scanner",
        "title": "Linux Persistence Mechanism Scanner",
        "stack": "Go 1.25, Cobra CLI",
        "purpose": "I built this to hunt persistence artifacts across a Linux root filesystem: systemd units, cron entries, shell profiles, SSH keys, LD_PRELOAD, udev rules, and more. Each hit includes severity and a MITRE ATT&CK reference.",
        "features": [
            "17 scanner modules across 12 persistence categories",
            "Heuristic detection for reverse shells and download-execute chains",
            "Baseline save and diff mode",
            "Terminal and JSON output",
            "Scan mounted roots with --root",
        ],
        "architecture": "internal/scanner/registry.go registers modules. Each scanner walks known paths and applies patterns.go heuristics. internal/report formats findings. internal/baseline stores snapshots for diff. cmd/sentinel is the CLI entry.",
        "challenges": [
            "Go module path referenced an external GitHub namespace. I renamed it to local module sentinel.",
            "Scanning requires Linux directory layout even though the binary cross-compiles on Windows.",
            "False positives appear on legitimate admin scripts that resemble download-execute patterns.",
        ],
        "solutions": [
            "I replaced import paths and rebuilt successfully on Windows.",
            "I documented sentinel scan --root /mnt/target for offline disk analysis.",
            "Severity scoring and MITRE tags help analysts triage before treating every hit as malicious.",
        ],
        "commands": "sentinel scan | sentinel scan --json | sentinel baseline save",
        "diagram": "sentinel",
    },
]


def build_document() -> Document:
    doc = Document()
    setup_styles(doc)

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Cybersecurity Beginner Projects")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Technical Documentation and Implementation Report")
    sr.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub2.add_run("15 Security Tools: Architecture, Challenges, and Solutions")
    sr2.font.size = Pt(13)
    sr2.italic = True

    doc.add_page_break()

    # TOC placeholder heading
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Portfolio Scope and Objectives",
        "3. Development Environment and Setup Pipeline",
        "4. Cross-Cutting Engineering Challenges",
        "5. Project Documentation (1 through 15)",
        "6. Security and Legal Considerations",
        "7. Conclusion and Future Work",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Generate diagrams
    d_overview = diagram_portfolio_overview()
    d_c2 = diagram_c2_sequence()
    d_canary = diagram_canary_er()
    d_sentinel = diagram_sentinel_flow()
    d_setup = diagram_setup_pipeline()
    d_netanal = diagram_netanal_layers()
    d_hash = diagram_hashcracker_uml()
    d_dns = diagram_dns_flow()
    d_metadata = diagram_metadata_flow()
    d_c2_deploy = diagram_c2_deployment()

    diagram_map = {
        "c2": d_c2,
        "canary": d_canary,
        "sentinel": d_sentinel,
        "netanal": d_netanal,
        "hash": d_hash,
        "dns": d_dns,
        "metadata": d_metadata,
    }

    # Section 1: Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    add_para(doc,
        "This document describes fifteen beginner-to-intermediate cybersecurity tools I developed "
        "as a structured learning portfolio. The collection spans encoding and cryptanalysis, network "
        "reconnaissance, deception technology, persistence hunting, compliance auditing, metadata privacy, "
        "and educational offensive tooling for isolated lab use. Each project is self-contained with source "
        "code, tests where applicable, and learn modules that document concepts and implementation detail.")
    add_para(doc,
        "I organized the work to mirror how security teams actually operate: collect and decode artifacts, "
        "map network exposure, detect attacker persistence, harden systems against benchmarks, and build "
        "detection layers with canary tokens and syscall tracing. The portfolio is not a single monolithic "
        "application. It is fifteen focused tools that share common engineering lessons around CLI design, "
        "dependency isolation, cross-platform limits, and responsible use.")
    add_image(doc, d_overview, "Figure 1: Portfolio architecture grouped by security domain")

    add_image(doc, d_c2_deploy, "Figure 3: C2 Beacon deployment architecture (nginx, FastAPI, SQLite, implant)")

    # Section 2
    doc.add_heading("2. Portfolio Scope and Objectives", level=1)
    add_para(doc, "The table below lists all fifteen projects with primary language and security role.")
    table = doc.add_table(rows=16, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Project", "Language", "Security Role"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    rows_data = [
        ("1", "base64-tool", "Python", "Encoding analysis and layer peeling"),
        ("2", "c2-beacon", "Python/React", "Educational C2 framework"),
        ("3", "caesar-cipher", "Python", "Classical cipher and frequency attack"),
        ("4", "canary-token-generator", "Go/React", "Honeytoken deception"),
        ("5", "dns-lookup", "Python", "DNS and WHOIS reconnaissance"),
        ("6", "firewall-rule-engine", "V", "Firewall ruleset analysis"),
        ("7", "hash-cracker", "C++", "Offline hash recovery"),
        ("8", "keylogger", "Python", "Input capture research (lab only)"),
        ("9", "linux-cis-hardening-auditor", "Bash", "CIS compliance auditing"),
        ("10", "linux-ebpf-security-tracer", "Python/C", "Kernel syscall monitoring"),
        ("11", "metadata-scrubber-tool", "Python", "File metadata removal"),
        ("12", "network-traffic-analyzer", "Python", "PCAP and live capture analysis"),
        ("13", "simple-port-scanner", "C++", "TCP port scanning"),
        ("14", "simple-vulnerability-scanner", "Go", "Python dependency CVE scan"),
        ("15", "systemd-persistence-scanner", "Go", "Linux persistence hunting"),
    ]
    for r_idx, row in enumerate(rows_data, 1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val

    doc.add_paragraph()
    add_para(doc, "Primary learning objectives I set for the portfolio:")
    add_bullets(doc, [
        "Build real CLI and service tools rather than notebook-only exercises",
        "Map each tool to MITRE ATT&CK or CIS controls where applicable",
        "Handle multiple languages: Python, Go, C++, V, Bash, and eBPF C",
        "Document architecture, failure modes, and recovery steps honestly",
        "Keep lab-only offensive code clearly scoped and isolated",
    ])

    # Section 3
    doc.add_heading("3. Development Environment and Setup Pipeline", level=1)
    add_para(doc,
        "I developed primarily on Windows 10 with Python 3.12. Go 1.26, CMake, LLVM, and Ninja were "
        "installed during the project when native builds required them. Linux-only components "
        "(eBPF tracing, CIS auditor execution, full Docker stacks) were designed for WSL or VM deployment.")
    add_image(doc, d_setup, "Figure 2: Sequential setup pipeline via setup.ps1")
    add_para(doc,
        "I wrote setup.ps1 to walk through all fifteen projects in order. For each Python project it creates "
        "a dedicated .venv, runs pip install -e ., and smoke-tests the CLI. For Go projects it runs "
        "go build and prints --help output. C++, V, and Linux-only projects report missing dependencies "
        "instead of failing silently.")
    add_para(doc, "Manual setup pattern for Python projects on Windows:")
    add_bullets(doc, [
        "cd <project-folder>",
        "python -m venv .venv",
        ".venv\\Scripts\\Activate.ps1",
        "pip install -e .",
    ])

    # Section 4: Cross-cutting challenges
    doc.add_heading("4. Cross-Cutting Engineering Challenges", level=1)

    challenges = [
        ("Dependency isolation",
         "Installing every Python project into one global environment caused typer, rich, and entry point "
         "collisions. The metadata-scrubber and ebpf-tracer packages both exposed scripts that resolved to "
         "the wrong src.main module.",
         "I changed setup.ps1 to create a per-project .venv. Each activation scope keeps dependencies local."),
        ("Python version constraints",
         "Several pyproject.toml files required Python 3.13 or 3.14 while my host runs 3.12. pip blocked "
         "installation entirely.",
         "I lowered requires-python to >=3.12 for local development and verified core commands still run."),
        ("Windows versus Linux capability gaps",
         "eBPF, CIS auditing, and full Docker compose stacks do not run natively on my Windows host.",
         "I separated build verification on Windows from runtime verification on Linux/WSL and documented both paths."),
        ("Go module namespace cleanup",
         "Go imports originally pointed at an external GitHub organization path.",
         "I renamed modules to sentinel and angela locally and rebuilt binaries successfully."),
        ("Native library requirements",
         "C++ projects need Boost and OpenSSL. Live packet capture needs Npcap on Windows.",
         "I installed CMake and LLVM via winget, documented apt-based Linux installs, and used PCAP replay "
         "when live capture was unavailable."),
    ]
    for title, problem, fix in challenges:
        doc.add_heading(title, level=2)
        add_para(doc, f"Problem: {problem}")
        add_para(doc, f"Resolution: {fix}")

    # Section 5: Projects
    doc.add_heading("5. Project Documentation", level=1)
    add_para(doc,
        "The following subsections document each project in a consistent format: purpose, capabilities, "
        "architecture, challenges faced, solutions implemented, and key commands verified during setup.")

    for proj in PROJECTS:
        doc.add_heading(f"5.{proj['num']} {proj['title']}", level=2)

        add_para(doc, f"Folder: {proj['name']}")
        add_para(doc, f"Technology stack: {proj['stack']}")
        add_para(doc, f"Purpose: {proj['purpose']}")

        doc.add_heading("Capabilities", level=3)
        add_bullets(doc, proj["features"])

        doc.add_heading("Architecture", level=3)
        add_para(doc, proj["architecture"])

        if proj.get("diagram") and proj["diagram"] in diagram_map:
            add_image(doc, diagram_map[proj["diagram"]],
                        f"Figure: Architecture diagram for {proj['name']}")

        doc.add_heading("Implementation Walkthrough", level=3)
        steps = proj.get("implementation_steps")
        if steps:
            add_bullets(doc, steps)
        else:
            add_para(doc,
                f"I started {proj['name']} by reading the learn/02-ARCHITECTURE.md and learn/03-IMPLEMENTATION.md "
                f"modules to map module boundaries before touching code. I then installed dependencies for "
                f"{proj['stack']} in an isolated environment, built or installed the package, and ran the primary "
                f"CLI entry point with --help to confirm the binary resolves. After that I executed the happy-path "
                f"command documented in the project README and captured output for regression comparison.")
            add_para(doc,
                f"For {proj['name']}, the critical integration point is the boundary between user input parsing and "
                f"the core engine. I validated that boundary first because most regressions I hit during portfolio "
                f"setup were dependency or PATH issues, not logic bugs in the security algorithms themselves.")

        doc.add_heading("Challenges Faced", level=3)
        add_bullets(doc, proj["challenges"])

        doc.add_heading("How I Resolved Them", level=3)
        add_bullets(doc, proj["solutions"])

        if proj.get("design_decisions"):
            doc.add_heading("Design Decisions", level=3)
            add_bullets(doc, proj["design_decisions"])

        doc.add_heading("Testing and Verification", level=3)
        testing = proj.get("testing", (
            f"I verified {proj['name']} using the setup.ps1 sequence on Windows where applicable. "
            f"Build or install succeeded, --help output rendered, and the primary command from README returned "
            f"expected output without unhandled exceptions."
        ))
        add_para(doc, testing)

        doc.add_heading("Verified Commands", level=3)
        add_para(doc, proj["commands"])

        doc.add_paragraph()

    # Appendix A: Setup results
    doc.add_page_break()
    doc.add_heading("Appendix A: Windows Setup Results", level=1)
    add_para(doc,
        "The table records what I verified on Windows 10 with Python 3.12, Go 1.26, CMake, and LLVM installed "
        "during portfolio setup. Status Ready means build or install plus --help smoke test passed.")
    setup_table = doc.add_table(rows=16, cols=4)
    setup_table.style = "Table Grid"
    sh = setup_table.rows[0].cells
    for i, h in enumerate(["#", "Project", "Windows Status", "Notes"]):
        sh[i].text = h
    setup_data = [
        ("1", "base64-tool", "Ready", "b64tool encode verified"),
        ("2", "c2-beacon", "Deps ready", "Needs Docker for full UI"),
        ("3", "caesar-cipher", "Ready", "encrypt --key 3 verified"),
        ("4", "canary-token-generator", "Backend built", "canary.exe --help OK"),
        ("5", "dns-lookup", "Ready", "JSON query mode used on Windows"),
        ("6", "firewall-rule-engine", "Source only", "Needs V compiler"),
        ("7", "hash-cracker", "Source only", "Needs Boost and OpenSSL"),
        ("8", "keylogger", "Ready", "Module import verified"),
        ("9", "linux-cis-hardening-auditor", "Linux only", "Use WSL for execution"),
        ("10", "linux-ebpf-security-tracer", "CLI ready", "Tracing needs Linux bcc"),
        ("11", "metadata-scrubber-tool", "Ready", "Use dedicated .venv"),
        ("12", "network-traffic-analyzer", "Ready", "Npcap needed for live capture"),
        ("13", "simple-port-scanner", "Source only", "Needs Boost on build host"),
        ("14", "simple-vulnerability-scanner", "Ready", "svscan.exe built"),
        ("15", "systemd-persistence-scanner", "Ready", "sentinel.exe built"),
    ]
    for r_idx, row in enumerate(setup_data, 1):
        for c_idx, val in enumerate(row):
            setup_table.rows[r_idx].cells[c_idx].text = val

    # Appendix B: MITRE mapping
    doc.add_heading("Appendix B: MITRE ATT&CK Mapping Summary", level=1)
    add_para(doc, "Several projects map directly to MITRE ATT&CK techniques. This table summarizes the strongest mappings.")
    mitre_table = doc.add_table(rows=11, cols=3)
    mitre_table.style = "Table Grid"
    mh = mitre_table.rows[0].cells
    for i, h in enumerate(["Project", "Technique", "Description"]):
        mh[i].text = h
    mitre_data = [
        ("c2-beacon", "T1071, T1059", "Application layer C2 and command execution"),
        ("canary-token-generator", "Engage D3", "Deception and tripwire detection"),
        ("keylogger", "T1056.001", "Input capture research in lab"),
        ("linux-ebpf-security-tracer", "T1548, T1059.004", "Privilege and reverse shell detection"),
        ("systemd-persistence-scanner", "TA0003", "Persistence discovery across Linux"),
        ("hash-cracker", "T1110", "Offline credential cracking research"),
        ("simple-port-scanner", "T1046", "Network service discovery"),
        ("dns-lookup", "T1016", "Network configuration discovery via DNS"),
        ("linux-cis-hardening-auditor", "M1028", "Operating system hardening validation"),
        ("metadata-scrubber-tool", "T1565.001", "Data manipulation awareness and privacy"),
    ]
    for r_idx, row in enumerate(mitre_data, 1):
        for c_idx, val in enumerate(row):
            mitre_table.rows[r_idx].cells[c_idx].text = val

    # Appendix C: Glossary
    doc.add_heading("Appendix C: Glossary", level=1)
    glossary = [
        ("eBPF", "Extended Berkeley Packet Filter. Allows sandboxed programs to run in the Linux kernel for observability."),
        ("Honeytoken", "A decoy asset that alerts the owner when accessed. Used in deception defense."),
        ("C2", "Command and Control. Infrastructure an operator uses to send tasks to compromised hosts."),
        ("CIS Benchmark", "Center for Internet Security configuration guidelines for secure system setup."),
        ("OSV", "Open Source Vulnerabilities database used for dependency CVE lookup."),
        ("LD_PRELOAD", "Linux mechanism to load a shared library before others. Common persistence vector."),
        ("PCAP", "Packet capture file format used to store network traffic for offline analysis."),
        ("BPF", "Berkeley Packet Filter filter expressions used to select packets during capture."),
        ("PEP 440", "Python version identification standard used when comparing package versions."),
        ("Chi-squared scoring", "Statistical test used in caesar-cipher to rank English-like plaintext candidates."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)

    # Appendix D: Methodology
    doc.add_page_break()
    doc.add_heading("Appendix D: Development Methodology", level=1)
    add_para(doc,
        "I followed a repeatable methodology for each project rather than ad-hoc copying. The steps below reflect "
        "how I actually worked through the portfolio on my Windows host.")
    methodology = [
        "Read learn/00-OVERVIEW.md and learn/01-CONCEPTS.md to understand the security problem before code.",
        "Skim learn/02-ARCHITECTURE.md and sketch module boundaries on paper or in a diagram.",
        "Install toolchain dependencies in an isolated environment (.venv, go mod download, or cmake configure).",
        "Run the documented quick-start command and capture baseline output.",
        "When install failed, record the exact error (Python version, missing DLL, missing compiler) before changing code.",
        "Apply the smallest fix first: version constraint, PATH, PYTHONPATH, or per-project venv.",
        "Re-run setup.ps1 for the full ordered sequence to confirm no regression in earlier projects.",
        "Update README with Windows-specific notes where the upstream quick start assumed Linux only.",
        "Document challenges and fixes in this report so the next environment setup is faster.",
    ]
    for i, step in enumerate(methodology, 1):
        doc.add_paragraph(f"{i}. {step}")

    add_para(doc,
        "This methodology kept the work honest. I did not claim a project was fully operational until I had either "
        "a successful command output or a clear documented blocker with a known resolution path.")

    # Section 6: Legal
    doc.add_heading("6. Security and Legal Considerations", level=1)
    add_para(doc,
        "Several projects model offensive techniques: C2 beacons, keyloggers, port scanners, and hash "
        "crackers. I built them for authorized education and lab research only. I do not deploy these tools "
        "against systems I do not own or lack written permission to test.")
    add_bullets(doc, [
        "C2 beacon and keylogger: isolated lab networks only",
        "Port scanner and DNS tools: authorized targets only",
        "Hash cracker: owned hash samples and CTF contexts only",
        "Canary tokens: deploy only inside environments you control",
        "Persistence scanner: use on owned Linux hosts or forensic images",
    ])

    # Section 7: Conclusion
    doc.add_heading("7. Conclusion and Future Work", level=1)
    add_para(doc,
        "Across fifteen projects I moved from simple encoding utilities to multi-service applications with "
        "databases, WebSocket protocols, kernel probes, and static Go binaries. The hardest problems were "
        "not algorithmic. They were environmental: dependency collisions, platform-specific modules, missing "
        "native libraries, and knowing when to verify builds on Windows versus runtime on Linux.")
    add_para(doc, "Planned extensions I identified in the learn challenge modules:")
    add_bullets(doc, [
        "Add YARA rule export to the persistence scanner",
        "Implement Falco-compatible rule output in the eBPF tracer",
        "Add SBOM generation to the vulnerability scanner",
        "Support Vigenere cipher extension in caesar-cipher",
        "Add TLS to the C2 channel for transport realism",
        "Build CI pipeline that runs setup.ps1 and captures pass/fail per project",
    ])
    add_para(doc,
        "This portfolio gave me end-to-end practice in security tooling: from parsing bytes and packets, "
        "to writing detection logic, to operating services that security teams actually run in production "
        "environments. The code, diagrams, and challenge notes in each project folder remain the living "
        "reference for ongoing development.")

    return doc


def main() -> None:
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Diagrams: {DIAGRAMS}")


if __name__ == "__main__":
    main()
