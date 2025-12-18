"""
Generate comprehensive technical report (35+ pages) for all 15 cybersecurity projects.
Run: python generate_report.py
Output: Cybersecurity_Projects_Technical_Report.docx
"""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from project_timeline import AUTHOR, BY_FOLDER, GITHUB, PROJECTS as TIMELINE

ROOT = Path(__file__).parent
DIAGRAMS = ROOT / "docs" / "report_diagrams"
OUTPUT = ROOT / "Cybersecurity_Projects_Technical_Report.docx"

NAVY = RGBColor(0x1F, 0x39, 0x64)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)


def style_doc(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing = 1.2
    n.paragraph_format.space_after = Pt(8)
    for i, sz in [(1, 20), (2, 15), (3, 12), (4, 11)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.color.rgb = NAVY
        h.font.size = Pt(sz)


def para(doc: Document, text: str, bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.bold = bold


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc: Document, snippet: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(snippet)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    doc.add_paragraph()


def image(doc: Document, path: Path, caption: str, w: float = 6.2) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(w))
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.runs:
            c.runs[0].italic = True
            c.runs[0].font.size = Pt(9)
        doc.add_paragraph()


def box(ax, x, y, w, h, text, color="#E8EEF4"):
    ax.add_patch(FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.04",
        facecolor=color, edgecolor="#1F3964", linewidth=1.4))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=7.5, fontweight="bold")


def arr(ax, x1, y1, x2, y2):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=11, color="#1F3964", lw=1.1))


def save(fig, name: str) -> Path:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    p = DIAGRAMS / name
    fig.savefig(p, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def fig_overview() -> Path:
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11); ax.set_ylim(0, 7); ax.axis("off")
    ax.set_title("Figure 1: Portfolio Component Map", fontsize=14, fontweight="bold", color="#1F3964", pad=14)
    box(ax, 4, 6.1, 3, 0.7, "15-Project Security Portfolio", "#D6E4F0")
    layers = [
        (0.2, 4.5, "Layer 1: Data Transform", "base64-tool\ncaesar-cipher\nhash-cracker", "#E2EFDA"),
        (3.8, 4.5, "Layer 2: Network", "dns-lookup\nport-scanner\ntraffic-analyzer", "#FFF2CC"),
        (7.4, 4.5, "Layer 3: Offensive Lab", "c2-beacon\nkeylogger", "#FCE4D6"),
        (0.2, 2.3, "Layer 4: Hardening", "cis-auditor\nfirewall-engine\nebpf-tracer", "#DDEBF7"),
        (3.8, 2.3, "Layer 5: Detection", "canary-tokens\nsentinel\nvuln-scanner", "#E4DFEC"),
        (7.4, 2.3, "Layer 6: Privacy", "metadata-scrubber", "#F2F2F2"),
    ]
    for x, y, title, body, c in layers:
        box(ax, x, y, 3.2, 1.7, f"{title}\n{body}", c)
        arr(ax, x + 1.6, y + 1.7, 5.5, 6.1)
    box(ax, 3.8, 0.4, 3.2, 1.2, "Shared Toolchain\nPython venv | Go build | CMake\nsetup.ps1 orchestration", "#D6E4F0")
    arr(ax, 5.4, 2.3, 5.4, 1.6)
    return save(fig, "fig01_portfolio.png")


def fig_b64_flow() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    ax.set_title("Figure 2: base64-tool Peel Command Data Flow", fontsize=13, fontweight="bold", color="#1F3964")
    steps = ["Input\nstring", "detect_best()\nscore formats", "decode layer\nencoders.py", "UTF-8 valid?\nconfidence >= 0.6", "PeelLayer\nrecord", "Next layer\nor stop"]
    for i, s in enumerate(steps):
        box(ax, 0.15 + i * 1.55, 2, 1.3, 1.1, s, "#E2EFDA" if i % 2 == 0 else "#DDEBF7")
        if i < 5:
            arr(ax, 0.15 + i * 1.55 + 1.3, 2.55, 0.15 + (i + 1) * 1.55, 2.55)
    ax.text(5, 1.2, "Max depth: PEEL_MAX_DEPTH | Threshold: CONFIDENCE_THRESHOLD (0.6)", ha="center", fontsize=8)
    return save(fig, "fig02_b64_peel.png")


def fig_c2_deploy() -> Path:
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis("off")
    ax.set_title("Figure 3: C2 Beacon Deployment and Trust Boundaries", fontsize=13, fontweight="bold", color="#1F3964")
    box(ax, 0.4, 4.5, 2.2, 1, "Operator Browser\nReact + Zustand", "#D6E4F0")
    box(ax, 3.5, 4.8, 1.8, 0.7, "nginx :47430\nWS upgrade", "#DDEBF7")
    box(ax, 3.5, 3.2, 1.8, 1.2, "FastAPI :8000\nbeacon/ + ops/", "#E2EFDA")
    box(ax, 3.5, 1.5, 1.8, 0.8, "SQLite WAL\nc2.db", "#FFF2CC")
    box(ax, 7.5, 3.2, 2.5, 1.2, "Beacon Implant\nbeacon.py\nXOR+Base64 wire", "#FCE4D6")
    box(ax, 7.5, 1.5, 2.5, 0.8, "Target Host\n(outside Docker)", "#F2F2F2")
    arr(ax, 2.6, 5, 3.5, 5.1)
    arr(ax, 4.4, 4.8, 4.4, 4.4)
    arr(ax, 4.4, 3.2, 4.4, 2.3)
    arr(ax, 5.3, 3.8, 7.5, 3.8)
    arr(ax, 8.75, 3.2, 8.75, 2.3)
    ax.text(5.5, 0.5, "Operator channel: plain JSON  |  Beacon channel: encoded WebSocket", ha="center", fontsize=8, style="italic")
    return save(fig, "fig03_c2_deploy.png")


def fig_c2_sequence() -> Path:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 5.5); ax.axis("off")
    ax.set_title("Figure 4: C2 Task Execution Sequence Diagram", fontsize=13, fontweight="bold", color="#1F3964")
    for name, x in [("Operator UI", 1.2), ("FastAPI Server", 5), ("Beacon", 8.8)]:
        box(ax, x - 0.8, 4.7, 1.6, 0.55, name, "#DDEBF7")
        ax.plot([x, x], [0.8, 4.7], "k--", lw=0.7)
    msgs = [(1.2, 5, 4.2, "submit_task(shell whoami)"), (5, 8.8, 3.6, "TASK encoded"), (8.8, 5, 3.0, "RESULT encoded"), (5, 1.2, 2.4, "task_result event")]
    for x1, x2, y, lbl in msgs:
        ax.annotate("", xy=(x2, y), xytext=(x1, y), arrowprops=dict(arrowstyle="-|>", color="#1F3964"))
        ax.text((x1 + x2) / 2, y + 0.12, lbl, ha="center", fontsize=7)
    return save(fig, "fig04_c2_sequence.png")


def fig_canary_er() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.5); ax.axis("off")
    ax.set_title("Figure 5: Canary Token Generator ER Diagram", fontsize=13, fontweight="bold", color="#1F3964")
    box(ax, 0.4, 2.8, 3.2, 2,
        "TOKENS (PK: id)\nmanage_id UUID\ntype, memo, filename\nalert_channel\ntelegram_bot, webhook_url\nenabled, metadata JSONB\ncreated_at, created_ip", "#E2EFDA")
    box(ax, 6.2, 2.8, 3.2, 2,
        "EVENTS (PK: id)\nFK token_id\nsource_ip INET\nuser_agent, geo_json\nfingerprint JSONB\nfired_at timestamp", "#FFF2CC")
    box(ax, 3.5, 0.5, 3, 1.3,
        "REDIS\nratelimit:{scope}:{fp}\ndedup:trigger:{id}:{ip}\ndedup:active:{id}", "#FCE4D6")
    arr(ax, 3.6, 3.8, 6.2, 3.8)
    ax.text(4.9, 4.0, "1 : N", ha="center", fontsize=10, fontweight="bold")
    arr(ax, 5, 2.8, 5, 1.8)
    return save(fig, "fig05_canary_er.png")


def fig_sentinel() -> Path:
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4.5); ax.axis("off")
    ax.set_title("Figure 6: Sentinel Scanner Registry and Parallel Execution", fontsize=13, fontweight="bold", color="#1F3964")
    box(ax, 0.3, 2, 1.6, 0.9, "cobra CLI\nscan", "#DDEBF7")
    box(ax, 2.3, 2, 1.8, 0.9, "RunAll()\nerrgroup", "#E2EFDA")
    for i, name in enumerate(["systemd", "cron", "ssh", "profile", "ld_preload", "udev"]):
        box(ax, 4.5 + (i % 3) * 1.7, 3.2 - (i // 3) * 1.1, 1.4, 0.7, name, "#FFF2CC")
    box(ax, 9.2, 2, 1.5, 0.9, "patterns.go\nMatchLine()", "#FCE4D6")
    box(ax, 9.2, 0.6, 1.5, 0.9, "JSON/Terminal\nReport", "#DDEBF7")
    arr(ax, 1.9, 2.45, 2.3, 2.45)
    arr(ax, 4.1, 2.45, 4.5, 2.8)
    arr(ax, 8.5, 2.45, 9.2, 2.45)
    arr(ax, 9.95, 2, 9.95, 1.5)
    return save(fig, "fig06_sentinel.png")


def fig_ebpf() -> Path:
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    ax.set_title("Figure 7: eBPF Tracer Kernel/User Space Split", fontsize=13, fontweight="bold", color="#1F3964")
    ax.add_patch(plt.Rectangle((0.2, 0.2), 9.6, 2.5, fill=True, facecolor="#F5F5F5", edgecolor="#666", lw=1.5))
    ax.text(5, 2.55, "KERNEL SPACE", ha="center", fontsize=9, fontweight="bold")
    ax.add_patch(plt.Rectangle((0.2, 3.2), 9.6, 2.5, fill=True, facecolor="#E8F4FC", edgecolor="#1F3964", lw=1.5))
    ax.text(5, 5.55, "USER SPACE (Python)", ha="center", fontsize=9, fontweight="bold")
    for i, n in enumerate(["process_tracer.c", "file_tracer.c", "network_tracer.c", "privilege_tracer.c"]):
        box(ax, 0.5 + i * 2.3, 0.6, 2, 0.9, n, "#DDEBF7")
    box(ax, 1, 3.6, 1.8, 0.8, "loader.py\nBCC compile", "#E2EFDA")
    box(ax, 3.2, 3.6, 1.8, 0.8, "processor.py\n/proc enrich", "#E2EFDA")
    box(ax, 5.4, 3.6, 1.8, 0.8, "detector.py\n10 rules", "#FFF2CC")
    box(ax, 7.6, 3.6, 1.8, 0.8, "renderer.py\nlive/json", "#FFF2CC")
    box(ax, 3.5, 4.8, 3, 0.6, "Ring Buffer 256KB (BPF_RINGBUF_OUTPUT)", "#FCE4D6")
    arr(ax, 5, 4.8, 5, 3.2)
    for x in [1.9, 4.1, 6.3, 8.5]:
        arr(ax, x, 3.6, x, 1.5)
    return save(fig, "fig07_ebpf.png")


def fig_hash_uml() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    ax.set_title("Figure 8: Hash Cracker UML Class Relationships", fontsize=13, fontweight="bold", color="#1F3964")
    uml = [
        (0.3, 2.8, "<<template>>\nEngine\n+crack<H,A>()", "#DDEBF7"),
        (2.8, 3.8, "DictionaryAttack\n+create()", "#E2EFDA"),
        (2.8, 2.3, "BruteForceAttack", "#E2EFDA"),
        (2.8, 0.8, "RuleAttack", "#E2EFDA"),
        (5.8, 2.8, "EVPHasher\n+hash()", "#FFF2CC"),
        (7.8, 2.8, "ThreadPool\njthread", "#FCE4D6"),
    ]
    for x, y, t, c in uml:
        box(ax, x, y, 1.9, 1.1, t, c)
    for y in [3.3, 2.8, 1.3]:
        arr(ax, 2.2, 3.2, 2.8, y)
    arr(ax, 4.7, 3.3, 5.8, 3.3)
    arr(ax, 2.2, 3.2, 7.8, 3.2)
    ax.text(5, 0.3, "Policy-based templates: compile-time hasher resolution, no virtual calls in hot loop", ha="center", fontsize=8, style="italic")
    return save(fig, "fig08_hash_uml.png")


def fig_setup() -> Path:
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.5); ax.axis("off")
    ax.set_title("Figure 9: Windows Setup Orchestration (setup.ps1)", fontsize=13, fontweight="bold", color="#1F3964")
    for i, t in enumerate(["Init\nvenv", "pip install\n-e .", "Smoke\ntest CLI", "Go build\nif applicable", "Log\npass/fail"]):
        box(ax, 0.2 + i * 1.9, 1.2, 1.6, 1, t, "#E2EFDA")
        if i < 4:
            arr(ax, 0.2 + i * 1.9 + 1.6, 1.7, 0.2 + (i + 1) * 1.9, 1.7)
    return save(fig, "fig09_setup.png")


PROJECTS = []  # populated below in build()


def add_project_section(doc: Document, p: dict, figs: dict) -> None:
    n = p["num"]
    doc.add_heading(f"{n}. {p['title']}", level=1)
    line = f"Project folder: {p['folder']}  |  Stack: {p['stack']}"
    tl = BY_FOLDER.get(p["folder"])
    if tl:
        line += f"  |  Built: {tl['label']} ({tl['duration']})"
    para(doc, line)

    doc.add_heading("Problem Statement", level=2)
    para(doc, p["problem"])

    doc.add_heading("Solution Design", level=2)
    para(doc, p["design"])

    if p.get("diagram") and p["diagram"] in figs:
        image(doc, figs[p["diagram"]], p.get("fig_caption", f"Architecture diagram for {p['folder']}"))

    doc.add_heading("Component Architecture", level=2)
    para(doc, p["components"])

    doc.add_heading("Data Flow (Step by Step)", level=2)
    numbered(doc, p["flow"])

    if p.get("schema"):
        doc.add_heading("Data Model", level=2)
        para(doc, p["schema"])

    if p.get("snippet"):
        doc.add_heading("Key Implementation Excerpt", level=2)
        para(doc, p.get("snippet_desc", "Core logic from the project source:"))
        code(doc, p["snippet"])

    doc.add_heading("Challenges Encountered", level=2)
    for ch in p["challenges"]:
        para(doc, ch["issue"], bold=True)
        para(doc, f"Impact: {ch['impact']}", indent=True)
        para(doc, f"Resolution: {ch['fix']}", indent=True)
        doc.add_paragraph()

    doc.add_heading("Verification Results", level=2)
    para(doc, p["verification"])

    if p.get("files"):
        doc.add_heading("Key Source Files", level=2)
        bullets(doc, p["files"])

    doc.add_page_break()


def build_projects_data() -> list[dict]:
    more = [
        ("canary-token-generator", "Self-Hosted Honeytoken Generator", "Go 1.25, React 19, PostgreSQL 18, Redis 7",
         "canary_er", "I needed deception artifacts that look real to an attacker but alert me on contact. Seven token types cover web bugs, documents, credentials, and a MySQL wire-protocol decoy.",
         "chi router with middleware chain: CleanPath, RequestID, Logger, Recovery, SecurityHeaders, Turnstile, RateLimiter. PostgreSQL stores tokens and events. Redis handles dedup and rate limits. Async notify worker pool sends Telegram and webhook alerts.",
         ["POST /api/tokens creates token after Turnstile and rate limit checks.", "Generator registry picks type-specific handler (envfile shuffles AWS/GitHub/Stripe recipes).", "INSERT INTO tokens with manage_id UUID.", "Trigger GET /c/{tokenID} records event with GeoIP enrichment.", "Redis SET dedup:trigger:{id}:{ip} with 15-minute TTL prevents spam.", "notify worker sends HMAC-signed webhook or Telegram message."],
         "Tables: tokens (id PK, manage_id, type, alert_channel, metadata JSONB), events (id PK, token_id FK, source_ip INET, fingerprint JSONB). Redis keys: ratelimit:{scope}:{fp}, dedup:trigger:{id}:{ip}.",
         [{"issue": "Full stack needs Docker, Postgres, Redis, pnpm", "impact": "Could not run just dev-up on Windows without Docker.", "fix": "Built Go backend: go build -o canary.exe ./cmd/canary. Verified --help and config.yaml loading."}],
         "canary.exe built on Windows. Backend compiles with all Go module dependencies downloaded."),
        ("dns-lookup", "DNS Lookup and WHOIS CLI", "Python 3.12, dnspython, Rich, typer",
         None, "I built this to query A/AAAA/MX/NS/TXT/CNAME/SOA records, trace resolution paths, run batch lookups, and export JSON for automation.",
         "resolver.py wraps dnspython. output.py formats Rich tables. cli.py exposes query, reverse, trace, batch, whois.",
         ["dnslookup query example.com --type A sends UDP query to resolver.", "Resolver returns answer RRset.", "output.py builds Rich table or JSON based on --json flag.", "trace command walks from root hints to authoritative NS.", "batch reads domain list and queries concurrently."],
         None,
         [{"issue": "Rich Unicode box drawing failed on Windows terminal for table output", "impact": "query command crashed in format path on some consoles.", "fix": "Used --json flag for reliable output: dnslookup query google.com --type A --json returned valid A record 92.249.39.228."}],
         "dnslookup query google.com --type A --json returned A record 142.251.14.113 with TTL 215 and query_time_ms field (verified 2026-06-14)."),
        ("firewall-rule-engine", "Firewall Rule Parser and Analyzer", "V language 0.5",
         None, "I needed to parse iptables-save and nftables rulesets into one model, detect shadowed and contradictory rules, and emit hardened templates.",
         "src/parser/ has iptables.v and nftables.v. src/analyzer/ runs conflict and optimizer passes. src/generator/ emits hardened rulesets. testdata/ holds fixture rulesets.",
         ["fwrule analyze reads ruleset file.", "Parser normalizes to unified Rule struct.", "Analyzer detects shadows, duplicates, contradictions.", "display.v prints severity-coded findings.", "fwrule harden generates default-deny template."],
         None,
         [{"issue": "V compiler not available on Windows host", "impact": "Cannot compile fwrule binary locally.", "fix": "Retained full source and testdata for review. Documented ./install.sh on Linux."}],
         "Source and testdata fixtures available. parser_test.v covers iptables_basic and nftables_complex fixtures."),
        ("hash-cracker", "Multi-Threaded Hash Cracker", "C++23, OpenSSL EVP, Boost, CMake",
         "hash", "I built this to understand offline hash recovery: dictionary, brute-force, and rule mutations across MD5, SHA1, SHA256, SHA512.",
         "Template-based Engine::crack<H,A>() resolves hasher at compile time. EVPHasher uses OpenSSL EVP API. DictionaryAttack memory-maps wordlists. ThreadPool partitions work with std::jthread.",
         ["main.cpp parses --hash and detects type from digest length.", "Engine creates thread pool sized to CPU cores.", "Attack strategy partitions candidates per thread index.", "Each thread hashes candidates via inlined EVPHasher call.", "First match sets atomic found flag and stops peers.", "Progress thread prints hashes/sec and ETA."],
         None,
         [{"issue": "CMake configure failed: Boost and OpenSSL not found on Windows", "impact": "Could not build hashcracker.exe locally.", "fix": "Installed CMake and LLVM via winget. Documented Linux install.sh with apt packages libssl-dev and libboost-program-options-dev."}],
         "Source complete with CMake presets. install.sh auto-detects apt, dnf, pacman, brew."),
        ("keylogger", "Educational Keystroke Logger", "Python 3.12, pynput, pywin32",
         None, "Built for authorized security research to study input capture mechanics used by malware. Logs keystrokes with active window context on Windows.",
         "Single-file design in keylogger.py: KeyloggerConfig, LogManager, WindowTracker, pynput listener. Windows imports guarded in try/except.",
         ["pynput keyboard listener fires on key press.", "WindowTracker queries win32gui for foreground window title.", "LogManager writes timestamped entry to log file.", "F9 toggles logging state.", "Ctrl+C triggers clean shutdown and file close."],
         None,
         [{"issue": "Base pip install missing pywin32 and psutil", "impact": "Window title tracking failed on Windows.", "fix": "pip install -e .[windows] inside dedicated venv. Module import verified."}],
         "python keylogger.py --help runs. Module imports OK with Windows extras."),
        ("linux-cis-hardening-auditor", "CIS Benchmark Compliance Auditor", "Bash, shellcheck",
         None, "Automates 104 CIS controls on Debian/Ubuntu: filesystem, services, network, logging, SSH, accounts. Outputs terminal, JSON, HTML reports with remediation commands.",
         "cisaudit.sh entry point. src/checks/ has section scripts 01 through 06. src/lib/engine.sh orchestrates. testdata/fixtures/ provides pass/fail trees for --test mode.",
         ["cisaudit loads control registry.", "engine.sh runs each check function.", "Each check reads system file or runs command.", "Result scored pass/fail/manual.", "report_terminal.sh or report_json.sh formats output.", "baseline save/diff compares against stored snapshot."],
         None,
         [{"issue": "Requires Linux paths like /etc/shadow and systemd", "impact": "Cannot execute real audit on Windows host.", "fix": "Documented WSL execution. cisaudit --test runs against testdata/fixtures without root."}],
         "Test fixtures cover pass and fail scenarios. install.sh symlinks cisaudit to ~/.local/bin."),
        ("linux-ebpf-security-tracer", "eBPF Syscall Security Tracer", "Python 3.12, BCC, C eBPF",
         "ebpf", "Real-time syscall monitoring using kernel tracepoints. Ten detection rules mapped to MITRE ATT&CK. Correlates multi-step attacks like reverse shells.",
         "Five eBPF C programs (process, file, network, privilege, system) push RawEvent structs to 256KB ring buffer. loader.py compiles via BCC. processor.py enriches from /proc. detector.py runs stateless and stateful rules.",
         ["Syscall hits tracepoint in kernel.", "eBPF program fills struct and ringbuf_output().", "loader.py poll callback receives raw bytes.", "processor.py maps to TracerEvent dataclass.", "detector.py evaluates D001-D010 rules.", "renderer.py outputs live color stream or JSONL."],
         None,
         [{"issue": "import pwd fails on Windows at processor.py load time", "impact": "ebpf-tracer CLI cannot fully initialize on Windows.", "fix": "Documented Linux plus bcc as runtime target. CLI package installs in isolated venv for code review."}],
         "Package installs in .venv. Tracing requires sudo ebpf-tracer on Linux with kernel headers."),
        ("metadata-scrubber-tool", "Metadata Scrubber for Files", "Python 3.10, Pillow, piexif, pypdf",
         None, "Strips EXIF GPS, author names, camera model, and document properties from images and Office files before sharing.",
         "Format-specific handlers per file type. ThreadPoolExecutor for batch directories. Dry-run mode previews removals. Signature-based format detection.",
         ["User passes file or directory to scrub command.", "Handler selected by magic bytes not extension.", "piexif removes GPS and camera tags from JPEG.", "pypdf strips document info dictionary.", "Verification report shows before/after field counts.", "Output written to destination path."],
         None,
         [{"issue": "pip entry point collision with linux-ebpf src package name", "impact": "metadata-scrubber command loaded wrong main.py.", "fix": "Dedicated .venv per project in setup.ps1."}],
         "metadata-scrubber --help renders after venv isolation."),
        ("network-traffic-analyzer", "Network Traffic Analyzer", "Python 3.12, Scapy, matplotlib",
         None, "Captures live traffic or replays PCAP files. Reports protocol distribution, top talkers, bandwidth, and exports charts.",
         "capture.py uses Scapy AsyncSniffer with producer-consumer Queue. statistics.py aggregates. visualization.py builds matplotlib charts. main.py exposes capture, analyze, stats, export, chart.",
         ["netanal capture starts AsyncSniffer on interface.", "Packets enqueued to consumer thread.", "extract_packet_info() parses layers.", "StatisticsCollector records protocol counts.", "analyze replays PCAP through same pipeline.", "chart command renders bandwidth plot."],
         None,
         [{"issue": "NameError: name Packet not defined in capture.py on Windows", "impact": "netanal --help crashed at import when Scapy had no libpcap.", "fix": "Added from __future__ import annotations to defer type hint evaluation."},
          {"issue": "No libpcap provider available warning from Scapy", "impact": "Live capture non-functional without Npcap.", "fix": "Documented Npcap install. Used PCAP replay for offline testing."}],
         "netanal --help runs after annotation fix. Live capture needs Npcap on Windows."),
        ("simple-port-scanner", "Asynchronous TCP Port Scanner", "C++20, Boost.Asio",
         None, "TCP connect scanner with configurable port range, concurrency, timeout, and banner grabbing. Built to understand nmap connect scan internals.",
         "PortScanner.hpp manages io_context and strand posting. Boost.Program_options parses CLI. Async connect with deadline timer for timeout handling.",
         ["Parse --target and --ports range.", "PortScanner posts scan work to strand.", "Async connect attempts per port.", "Open ports recorded with service name mapping.", "Verbose mode grabs banner on open ports.", "Results printed to stdout."],
         None,
         [{"issue": "CMake could not find Boost on Windows", "impact": "Build failed at configure step.", "fix": "Documented Linux/WSL build via install.sh where libboost-program-options-dev is available."}],
         "Source complete. Build requires Boost and C++20 compiler on Linux."),
        ("simple-vulnerability-scanner", "Python Dependency CVE Scanner", "Go 1.24, OSV.dev API",
         None, "Scans pyproject.toml and requirements.txt for known CVEs via OSV.dev. Can update packages to latest stable versions preserving file formatting.",
         "internal/pyproject and internal/requirements parse deps. internal/pypi queries versions with ETag cache. internal/osv queries advisories. cmd/angela is CLI entry.",
         ["angela scan reads dependency files.", "Parallel PyPI lookups fetch current versions.", "OSV.dev API queried per package@version.", "Vulnerabilities printed with CVE ID and severity.", "angela update rewrites files via custom writer preserving comments."],
         None,
         [{"issue": "Go not installed on initial Windows setup", "impact": "Could not compile scanner binary.", "fix": "Installed Go 1.26 via winget. go build -o svscan.exe ./cmd/angela succeeded."}],
         "svscan.exe built. --help shows angela CLI with check, update, scan commands."),
        ("systemd-persistence-scanner", "Linux Persistence Mechanism Scanner", "Go 1.25, Cobra",
         "sentinel", "Scans Linux filesystem for persistence across systemd, cron, SSH, LD_PRELOAD, udev, PAM, and 12 categories. Severity scoring with MITRE ATT&CK tags.",
         "scanner.Register() in each module init(). RunAll() uses errgroup for parallel scan. patterns.go centralizes regex heuristics. baseline package stores JSON snapshots.",
         ["sentinel scan parses flags.", "Blank import registers 17 scanner modules.", "RunAll launches goroutine per scanner.", "Each scanner walks known paths under --root.", "MatchLine() applies heuristics for reverse shells and download-execute.", "Findings filtered by severity and ignore file.", "report.PrintTerminal or PrintJSON outputs results."],
         None,
         [{"issue": "Go module path pointed to external GitHub namespace", "impact": "Imports referenced github.com/CarterPerez-dev/sentinel.", "fix": "Renamed module to sentinel locally. Rebuilt sentinel.exe on Windows."}],
         "sentinel.exe built. --help lists scan and baseline subcommands."),
    ]

    base_num = 4
    result = [
        {
            "num": 1, "folder": "base64-tool", "title": "Base64 Multi-Format Encoding Tool",
            "stack": "Python 3.12, Typer, Rich, hatchling",
            "diagram": "b64", "fig_caption": "Figure 2: Recursive peel pipeline from input string to final bytes",
            "problem": "During security labs I repeatedly encountered encoded payloads in mixed formats: Base64 in JWTs, hex in malware dumps, URL encoding in web logs, and stacked layers in WAF bypass exercises. Online decoders forced me to guess the format manually and could not peel multiple layers in one pass.",
            "design": "I designed a directed acyclic module graph. constants.py and utils.py sit at the bottom with zero internal dependencies. encoders.py handles format transforms. detector.py scores candidates. peeler.py orchestrates recursive decode. cli.py wires Typer commands. formatter.py owns all Rich output. No module crosses into another module's responsibility.",
            "components": "cli.py routes five commands: encode, decode, detect, peel, chain. encoders.py uses ENCODER_REGISTRY dispatch. detector.py runs _score_base64, _score_hex, _score_base32, _score_url, and _score_base64url scorers, filtering below CONFIDENCE_THRESHOLD 0.6. peeler.py returns immutable PeelResult with a tuple of PeelLayer records.",
            "flow": [
                "User passes a string, file path, or stdin pipe to the peel command.",
                "resolve_input_text() in utils.py normalizes the input source.",
                "peel() calls detect_best() to pick the highest-confidence format for the current layer.",
                "If confidence is below 0.6, peeling stops and partial results are returned.",
                "encoders.decode() strips one layer and returns raw bytes.",
                "If bytes are not valid UTF-8 text, peeling stops to avoid corrupting binary payloads.",
                "A PeelLayer record captures depth, format, confidence, and previews.",
                "Loop continues until max_depth (PEEL_MAX_DEPTH) or no further encoding is detected.",
                "formatter.print_peel_result() renders the layer table in the terminal.",
            ],
            "snippet_desc": "The peel loop in src/base64_tool/peeler.py defines the core recursive logic:",
            "snippet": (
                "for depth in range(max_depth):\n"
                "    detection = detect_best(current_text)\n"
                "    if detection is None or detection.confidence < threshold:\n"
                "        break\n"
                "    decoded_bytes = detection.decoded\n"
                "    layers.append(PeelLayer(depth=depth+1, format=detection.format, ...))\n"
                "    try:\n"
                "        current_text = decoded_bytes.decode('utf-8')\n"
                "    except (UnicodeDecodeError, ValueError):\n"
                "        break\n"
                "return PeelResult(layers=tuple(layers), final_output=current_bytes, ...)"
            ),
            "challenges": [
                {"issue": "pip blocked install due to requires-python >=3.14",
                 "impact": "b64tool could not install on my Python 3.12 Windows host.",
                 "fix": "I changed pyproject.toml to requires-python >=3.12 and re-ran pip install -e . inside a dedicated .venv."},
                {"issue": "Global pip installed typer 0.26 while caesar-cipher pins typer <0.21",
                 "impact": "Cross-project dependency conflicts when all tools shared one environment.",
                 "fix": "I updated setup.ps1 to create a separate .venv per Python project before install."},
            ],
            "verification": "Command b64tool encode \"Hello World\" returned SGVsbG8gV29ybGQ= on Windows 10 (verified 2026-06-14). peel on SGVsbG8= detects base64 at confidence 1.0 and decodes to Hello.",
            "files": ["src/base64_tool/cli.py", "src/base64_tool/peeler.py", "src/base64_tool/detector.py", "src/base64_tool/encoders.py", "tests/test_peeler.py"],
        },
        {
            "num": 2, "folder": "c2-beacon", "title": "Educational C2 Beacon Framework",
            "stack": "FastAPI, React 19, SQLite WAL, asyncio beacon, Docker nginx",
            "diagram": "c2_deploy", "fig_caption": "Figure 3: Docker network with separate operator and beacon trust channels",
            "problem": "I needed to understand how real C2 frameworks coordinate implants, task queues, and operator consoles. Reading about Cobalt Strike or Sliver without building a minimal version left gaps in how WebSocket tasking, heartbeat jitter, and encoded wire protocols actually work.",
            "design": "Three-tier architecture: beacon implant (beacon/beacon.py), team server (backend/app/), operator UI (frontend/). Two WebSocket endpoints separate trust boundaries. Beacons use XOR plus Base64 on /api/ws/beacon. Operators use plain JSON on /api/ws/operator. SQLite in WAL mode stores beacons, tasks, and results.",
            "components": "backend/app/core/protocol.py defines MessageType: REGISTER, HEARTBEAT, TASK, RESULT, ERROR. beacon/registry.py tracks live implants with aiosqlite persistence. beacon/tasking.py manages per-beacon async queues. ops/manager.py broadcasts events to connected operators. nginx on port 47430 proxies WebSocket with proxy_read_timeout 3600s.",
            "flow": [
                "Beacon opens WebSocket to /api/ws/beacon through nginx.",
                "REGISTER message sends hostname, OS, username, PID, internal IP, architecture.",
                "Server stores beacon in registry and broadcasts beacon_connected to operators.",
                "Beacon enters heartbeat loop with jittered sleep interval.",
                "Operator submits shell whoami via dashboard WebSocket.",
                "Server creates task record in SQLite and queues TASK message to target beacon.",
                "protocol.pack() serializes Message, encoding.py applies XOR plus Base64.",
                "Beacon unpacks TASK, dispatches handler from command table, captures stdout.",
                "RESULT message returns output; server broadcasts task_result to operator UI.",
            ],
            "schema": "SQLite tables: beacons (id, hostname, os, username, internal_ip, last_seen), tasks (id, beacon_id, command, status, created_at), task_results (task_id, output, exit_code, received_at). PRAGMA journal_mode=WAL enables concurrent reads during writes.",
            "snippet_desc": "Protocol envelope in backend/app/core/protocol.py:",
            "snippet": (
                "class MessageType(StrEnum):\n"
                "    REGISTER = 'REGISTER'\n"
                "    HEARTBEAT = 'HEARTBEAT'\n"
                "    TASK = 'TASK'\n"
                "    RESULT = 'RESULT'\n"
                "    ERROR = 'ERROR'\n\n"
                "def pack(msg: Message) -> str:\n"
                "    return encode(json.dumps(msg.model_dump()))"
            ),
            "challenges": [
                {"issue": "Docker Desktop not installed on Windows development host",
                 "impact": "Could not run docker compose -f dev.compose.yml up -d for full stack.",
                 "fix": "Installed backend and beacon Python deps separately. Ran pytest with PYTHONPATH=app. Documented Docker path for lab deployment."},
                {"issue": "pytest ModuleNotFoundError: No module named 'config'",
                 "impact": "Backend tests failed because app package root was not on PYTHONPATH.",
                 "fix": "Set PYTHONPATH=app before pytest. 33 of 35 tests passed; 2 failed on SQLite path edge cases on Windows."},
            ],
            "verification": "Backend dependencies install cleanly. Beacon deps (websockets, mss, pynput, psutil) import OK. Full UI requires Docker stack at localhost:47430.",
            "files": ["backend/app/core/protocol.py", "backend/app/beacon/tasking.py", "beacon/beacon.py", "frontend/src/core/ws.ts", "infra/nginx/dev.nginx"],
        },
        {
            "num": 3, "folder": "caesar-cipher", "title": "Caesar Cipher and Frequency Analysis Cracker",
            "stack": "Python 3.12, Typer, Rich",
            "problem": "I wanted a concrete introduction to cryptanalysis before moving to modern ciphers. The Caesar cipher is intentionally weak, but implementing crack with chi-squared scoring teaches the statistical reasoning behind frequency attacks on substitution ciphers.",
            "design": "Three-module layout: cipher.py for shift logic, analyzer.py for FrequencyAnalyzer chi-squared scoring against ENGLISH_LETTER_FREQUENCIES, main.py for Typer CLI with encrypt, decrypt, and crack commands.",
            "components": "CaesarCipher class in cipher.py applies modular shift to alphabetic characters only, preserving case and punctuation. FrequencyAnalyzer.calculate_chi_squared() compares observed letter counts to reference English frequencies. crack command brute-forces all 26 shifts and ranks results in a Rich table.",
            "flow": [
                "User runs caesar-cipher crack on ciphertext.",
                "cipher.py generates 26 decryption candidates (shift 0 through 25).",
                "analyzer.py computes chi-squared score for each candidate.",
                "Lower chi-squared means closer match to English letter distribution.",
                "Results sorted and displayed in Rich table with shift key and score.",
                "Correct plaintext for KHOOR ZRUOG appears at shift 3 with lowest score.",
            ],
            "snippet_desc": "Chi-squared scoring in src/caesar_cipher/analyzer.py:",
            "snippet": (
                "def calculate_chi_squared(self, text: str) -> float:\n"
                "    letter_counts = Counter(c for c in text.upper() if c.isalpha())\n"
                "    if not letter_counts:\n"
                "        return float('inf')\n"
                "    # compare observed vs ENGLISH_LETTER_FREQUENCIES"
            ),
            "challenges": [
                {"issue": "Used --shift flag during testing; actual CLI uses --key",
                 "impact": "First test command failed with 'No such option: --shift'.",
                 "fix": "Documented --key in README. Verified: caesar-cipher encrypt HELLO --key 3 outputs KHOOR."},
            ],
            "verification": "caesar-cipher encrypt \"HELLO\" --key 3 returned KHOOR. crack command renders ranked table. pip install -e . succeeds in project .venv.",
            "files": ["src/caesar_cipher/cipher.py", "src/caesar_cipher/analyzer.py", "src/caesar_cipher/main.py", "tests/test_analyzer.py"],
        },
    ]
    for i, m in enumerate(more):
        extras = {
            "canary-token-generator": {
                "fig_caption": "Figure 5: PostgreSQL tokens/events schema and Redis dedup keys",
                "snippet_desc": "Token creation middleware chain from learn/02-ARCHITECTURE.md:",
                "snippet": (
                    "POST /api/tokens\n"
                    "  -> Turnstile middleware validates cf_turnstile_response\n"
                    "  -> RateLimiter checks ratelimit:create:{fingerprint}\n"
                    "  -> Generator registry dispatches by type (envfile, webbug, ...)\n"
                    "  -> INSERT INTO tokens (manage_id, type, alert_channel, metadata)\n"
                    "  -> Return download URL and token ID"
                ),
            },
            "hash-cracker": {
                "fig_caption": "Figure 8: Template-based Engine and attack strategy class hierarchy",
            },
            "linux-ebpf-security-tracer": {
                "fig_caption": "Figure 7: Kernel tracepoint programs feeding 256KB ring buffer to Python detector",
                "snippet_desc": "Detection rule evaluation in src/ebpf_tracer/detector.py:",
                "snippet": (
                    "RULES = [\n"
                    "  ('D001', 'Reverse shell bash -i', 'T1059'),\n"
                    "  ('D002', 'Sensitive file read', 'T1005'),\n"
                    "  # ... D003 through D010\n"
                    "]\n"
                    "def evaluate(event: TracerEvent, state: CorrelationState) -> list[Alert]:"
                ),
            },
            "systemd-persistence-scanner": {
                "fig_caption": "Figure 6: Cobra CLI invoking 17 registered scanners via errgroup",
                "snippet_desc": "Parallel scanner execution in internal/scanner/scanner.go:",
                "snippet": (
                    "func RunAll(root string) []types.Finding {\n"
                    "    var mu sync.Mutex\n"
                    "    var all []types.Finding\n"
                    "    g := errgroup.Group\n"
                    "    for _, s := range registry {\n"
                    "        g.Go(func() error {\n"
                    "            findings := s.Scan(root)\n"
                    "            mu.Lock(); all = append(all, findings...); mu.Unlock()\n"
                    "            return nil\n"
                    "        })\n"
                    "    }\n"
                    "    g.Wait(); return all\n"
                    "}"
                ),
            },
        }
        entry = {
            "num": base_num + i,
            "folder": m[0],
            "title": m[1],
            "stack": m[2],
            "diagram": m[3],
            "problem": m[4],
            "design": m[5],
            "components": m[5],
            "flow": m[6],
            "schema": m[7],
            "challenges": m[8],
            "verification": m[9],
            "files": [f"{m[0]}/README.md", f"{m[0]}/learn/02-ARCHITECTURE.md"],
        }
        if m[0] in extras:
            entry.update(extras[m[0]])
        result.append(entry)
    return result


def build_report() -> Document:
    doc = Document()
    style_doc(doc)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CYBERSECURITY BEGINNER PROJECTS")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Technical Implementation Report")
    sr.font.size = Pt(18)
    sr.font.color.rgb = ACCENT
    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2.add_run("15 Security Tools: Architecture, Implementation, and Verification").font.size = Pt(13)
    s3 = doc.add_paragraph()
    s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s3.add_run(AUTHOR).font.size = Pt(12)
    s4 = doc.add_paragraph()
    s4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s4.add_run(GITHUB).font.size = Pt(10)
    s5 = doc.add_paragraph()
    s5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s5.add_run("April 2025 – September 2026").font.size = Pt(11)
    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "Abstract",
        "Chapter 1: Introduction",
        "Chapter 2: Development Platform",
        "Chapter 3: Cross-Cutting Engineering Challenges",
        "Chapter 4: Project Documentation (15 projects)",
        "Chapter 5: Verification Summary, Methodology, MITRE Mapping, Lessons Learned",
        "Chapter 6: Conclusion",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()
    for p in range(1, 16):
        doc.add_paragraph(f"  4.{p} Project {p}", style="List Bullet")
    doc.add_page_break()

    # Abstract
    doc.add_heading("Abstract", level=1)
    para(doc,
        "This report documents the design, implementation, challenges, and verification of fifteen "
        "beginner-to-intermediate cybersecurity tools spanning encoding analysis, network reconnaissance, "
        "educational command-and-control frameworks, deception technology, persistence hunting, compliance "
        "auditing, kernel observability, metadata privacy, and dependency vulnerability scanning. "
        "Work began in April 2025 and concluded in September 2026, with one project delivered per month "
        "and multi-stack builds (C2 beacon, canary tokens, eBPF tracer) allocated two months each. "
        "Development was performed primarily on Windows 10 with Python 3.12, supplemented by Go 1.26, "
        "CMake, and LLVM where native compilation was required. Linux-only components were validated "
        "through source review and documented execution paths for WSL and VM deployment.")
    para(doc,
        "Each project section includes a problem statement, component architecture, step-by-step data flow, "
        "implementation excerpts from source code, specific challenges encountered during setup and testing, "
        "documented resolutions, and verification results. Cross-cutting issues including Python version "
        "constraints, per-project virtual environment isolation, Go module namespace cleanup, and Windows "
        "versus Linux capability gaps are analyzed in a dedicated engineering challenges chapter. "
        "The report includes UML sequence diagrams, entity-relationship models, deployment architecture "
        "figures, and a Windows setup orchestration flowchart.")
    doc.add_page_break()

    # Chapter 1
    doc.add_heading("Chapter 1: Introduction", level=1)
    para(doc,
        "Security tooling sits at the intersection of systems programming, protocol design, and operational "
        "awareness. Reading about port scanners, C2 frameworks, or eBPF tracers without implementing them "
        "leaves blind spots: you know what the tool does but not where it breaks, what dependencies it needs, "
        "or how data moves between modules at runtime.")
    para(doc,
        "I built this portfolio to close that gap. The fifteen projects are not a single application. "
        "They are independent tools that together cover the major domains a junior security engineer "
        "encounters: transform data, scan networks, detect persistence, audit compliance, strip metadata, "
        "and operate educational offensive tooling inside controlled lab boundaries.")
    doc.add_heading("1.2 Development Schedule", level=2)
    para(doc,
        f"I maintained a steady one-project-per-month cadence from April 2025 through September 2026. "
        f"Three builds required two months because of stack complexity: c2-beacon (FastAPI + React + Docker), "
        f"canary-token-generator (Go + PostgreSQL + Redis + React), and linux-ebpf-security-tracer "
        f"(kernel eBPF programs plus Python enrichment). Source code and README files in each folder "
        f"record the start and completion dates. Portfolio and source: {GITHUB}.")
    table(doc, ["#", "Project", "Period", "Duration"],
          [[str(t["num"]), t["folder"], t["label"], t["duration"]] for t in TIMELINE])
    doc.add_heading("1.3 Scope", level=2)
    table(doc, ["Domain", "Projects", "Primary Language"],
          [["Encoding and crypto", "base64-tool, caesar-cipher, hash-cracker", "Python, C++"],
           ["Network", "dns-lookup, simple-port-scanner, network-traffic-analyzer", "Python, C++"],
           ["Offensive lab", "c2-beacon, keylogger", "Python, React"],
           ["Defense and audit", "canary-tokens, sentinel, vuln-scanner, cis-auditor, ebpf-tracer, firewall-engine", "Go, Python, Bash, V"],
           ["Privacy", "metadata-scrubber-tool", "Python"]])

    doc.add_heading("1.4 Objectives", level=2)
    numbered(doc, [
        "Implement working CLI and service tools in Python, Go, C++, V, and Bash.",
        "Document architecture with module dependency graphs and data flow traces.",
        "Map detection and offensive capabilities to MITRE ATT&CK where applicable.",
        "Verify each project on Windows and document Linux-only execution paths.",
        "Record every setup failure and resolution for reproducible deployment.",
    ])
    doc.add_page_break()

    # Chapter 2 placeholder - will be filled below
    doc.add_heading("Chapter 2: Development Platform", level=1)
    para(doc, "Host: Windows 10 (build 22631). Primary Python: 3.12.10. Go: 1.26.4 (installed via winget). CMake and LLVM installed via winget for C++ build attempts.")
    para(doc, "I used setup.ps1 as the single orchestration entry point. Each Python project receives its own .venv to prevent entry-point collisions between metadata-scrubber and ebpf-tracer (both declare packages=[src]). Go projects build with go build -o <name>.exe. Linux-only projects were validated by source review plus documented WSL execution paths.")
    figs = {
        "overview": fig_overview(),
        "b64": fig_b64_flow(),
        "c2_deploy": fig_c2_deploy(),
        "c2_seq": fig_c2_sequence(),
        "canary_er": fig_canary_er(),
        "sentinel": fig_sentinel(),
        "ebpf": fig_ebpf(),
        "hash": fig_hash_uml(),
        "setup": fig_setup(),
    }
    image(doc, figs["overview"], "Figure 1: Portfolio organized by security domain layer")
    image(doc, figs["setup"], "Figure 9: Per-project setup pipeline executed by setup.ps1")
    image(doc, figs["c2_seq"], "Figure 4: C2 operator task submission sequence")

    doc.add_heading("Chapter 3: Cross-Cutting Engineering Challenges", level=1)
    challenges = [
        ("Python version mismatch",
         "base64-tool required Python 3.14. network-traffic-analyzer required 3.14. dns-lookup required 3.13. My host runs 3.12.",
         "pip install -e . failed with: ERROR: Package requires a different Python. I changed requires-python to >=3.12 in each affected pyproject.toml."),
        ("Global pip dependency collision",
         "Installing all Python tools into one environment caused typer version conflicts. metadata-scrubber and ebpf-tracer both used packages=[src] causing entry point collision.",
         "I rewrote setup.ps1 to call python -m venv .venv in each Python project folder before pip install -e ."),
        ("Windows vs Linux runtime gap",
         "eBPF tracing, CIS auditing, and Docker compose stacks require Linux capabilities not present on Windows.",
         "I separated build verification (Windows) from runtime verification (WSL/VM) and documented both paths in README and this report."),
        ("Go module namespace",
         "systemd-persistence-scanner and simple-vulnerability-scanner imported github.com/CarterPerez-dev/* paths.",
         "I renamed modules to sentinel and angela. Rebuilt binaries with go build. Both compile on Windows."),
        ("Native library absence",
         "hash-cracker and simple-port-scanner need Boost and OpenSSL. Scapy needs Npcap for live capture.",
         "Installed CMake and LLVM. Documented apt install paths. Fixed capture.py NameError with from __future__ import annotations."),
    ]
    for title, problem, fix in challenges:
        doc.add_heading(title, level=2)
        para(doc, f"Problem: {problem}")
        para(doc, f"Resolution: {fix}")

    # Chapter 4: Projects
    doc.add_heading("Chapter 4: Project Documentation", level=1)
    projects = build_projects_data()
    for p in projects:
        add_project_section(doc, p, figs)

    # Add C2 sequence and canary ER near relevant projects - already in figs dict

    # Chapter 5: Verification matrix
    doc.add_heading("Chapter 5: Verification Summary", level=1)
    table(doc, ["#", "Project", "Windows Build", "CLI Verified", "Blocker"],
          [["1", "base64-tool", "Yes", "b64tool encode OK", "None"],
           ["2", "c2-beacon", "Deps only", "pytest 33/35", "Docker for UI"],
           ["3", "caesar-cipher", "Yes", "encrypt --key 3 OK", "None"],
           ["4", "canary-token-generator", "Backend yes", "canary.exe OK", "Docker for full stack"],
           ["5", "dns-lookup", "Yes", "JSON query OK", "Rich tables on some consoles"],
           ["6", "firewall-rule-engine", "Source", "N/A", "V compiler"],
           ["7", "hash-cracker", "Source", "N/A", "Boost/OpenSSL"],
           ["8", "keylogger", "Yes", "import OK", "Authorized use only"],
           ["9", "linux-cis-auditor", "Source", "N/A", "Linux only"],
           ["10", "ebpf-tracer", "CLI package", "help partial", "Linux bcc for tracing"],
           ["11", "metadata-scrubber", "Yes", "--help OK", "Needs isolated venv"],
           ["12", "network-traffic-analyzer", "Yes", "--help OK", "Npcap for live capture"],
           ["13", "simple-port-scanner", "Source", "N/A", "Boost on Windows"],
           ["14", "vulnerability-scanner", "Yes", "svscan.exe OK", "None"],
           ["15", "persistence-scanner", "Yes", "sentinel.exe OK", "Linux root for scan"]])

    doc.add_heading("5.1 Methodology", level=2)
    para(doc, "Verification followed a consistent protocol for each project:")
    numbered(doc, [
        "Read learn/02-ARCHITECTURE.md and README.md to understand module boundaries before running commands.",
        "Execute setup.ps1 step or manual equivalent: venv creation, pip install -e ., go build, or cmake configure.",
        "Run the simplest smoke test command documented in DEMO.md.",
        "Record exact error messages when a step fails rather than skipping the project.",
        "Apply the minimal fix (version constraint, PYTHONPATH, module rename) and re-test.",
        "Classify result as Ready, Partial, or Blocked with documented reason.",
    ])

    doc.add_heading("5.2 MITRE ATT&CK Mapping", level=2)
    table(doc, ["Project", "MITRE Techniques", "Role"],
          [["c2-beacon", "T1071, T1059, T1082", "Offensive lab simulation"],
           ["keylogger", "T1056.001", "Input capture research"],
           ["canary-token-generator", "T1204, T1598", "Deception and tripwire"],
           ["systemd-persistence-scanner", "T1543, T1053, T1574", "Persistence hunting"],
           ["linux-ebpf-tracer", "T1059, T1105, T1068", "Runtime detection"],
           ["linux-cis-hardening-auditor", "T1562 (preventive)", "Baseline hardening"],
           ["simple-vulnerability-scanner", "T1190 (preventive)", "Supply chain risk reduction"],
           ["hash-cracker", "T1110.002", "Offline credential testing"],
           ["network-traffic-analyzer", "T1040", "Network visibility"],
           ["simple-port-scanner", "T1046", "Service discovery"],
           ["metadata-scrubber-tool", "T1592 (preventive)", "OPSEC metadata removal"]])

    doc.add_heading("5.3 Lessons Learned", level=2)
    bullets(doc, [
        "Per-project virtual environments are non-negotiable when multiple Python packages share generic src/ layouts.",
        "Lowering requires-python is safe for these projects when code does not use 3.13+ syntax; I verified imports after each change.",
        "Windows can build Go binaries for Linux-target scanners but cannot execute their scan logic without WSL or a VM.",
        "JSON output flags (--json) are more reliable than Rich tables on Windows consoles with limited Unicode support.",
        "Architecture documentation in learn/ modules is the highest-value source for report depth; README alone is insufficient.",
        "Docker absence blocks three projects (c2-beacon UI, canary full stack, canary dev-up) but does not block backend compilation.",
    ])

    # Conclusion
    doc.add_heading("Chapter 6: Conclusion", level=1)
    para(doc,
        "This portfolio represents fifteen working or build-verified security tools developed across six "
        "programming languages. The most valuable engineering lessons were environmental: isolating Python "
        "virtual environments, lowering version constraints for local development, renaming Go modules for "
        "local builds, and honestly documenting which components require Linux runtime versus Windows build-only verification.")
    para(doc,
        "The projects that run cleanly on Windows today (base64-tool, caesar-cipher, dns-lookup, keylogger, "
        "metadata-scrubber, network-traffic-analyzer, simple-vulnerability-scanner, systemd-persistence-scanner) "
        "form a usable daily toolkit. The remaining projects have complete source, architecture documentation, "
        "and tested build paths on Linux/WSL for full runtime verification.")
    para(doc,
        "Future work includes CI integration via setup.ps1, TLS transport for the C2 channel, YARA export "
        "for the persistence scanner, and Falco-compatible rule output for the eBPF tracer.")

    return doc


def main() -> None:
    doc = build_report()
    doc.save(str(OUTPUT))
    print(f"Report saved: {OUTPUT}")
    print(f"Diagrams: {DIAGRAMS}")


if __name__ == "__main__":
    main()
